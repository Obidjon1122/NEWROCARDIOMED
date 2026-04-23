import os
import io
from typing import Dict, Any
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Path to original DOCX template files
TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', '..', 'protocoles')

TEMPLATE_FILES: Dict[int, str] = {
    1:  'Печень нов..docx',
    2:  'Почки нов..docx',
    3:  'Щитовидная железа..docx',
    4:  'Поджелудочная новый.docx',
    5:  'Селезёнка новый.docx',
    6:  'Простата новый.docx',
    7:  'Молочные железы новый.docx',
    8:  'Мочевой пузырь новый..docx',
    9:  'Малый таз новый.docx',
    10: 'Надпочечники новый.docx',
    11: 'Коленный сустав новый.docx',
    12: 'Первый ориместр новый.docx',
    13: 'Плод новый.docx',
    14: 'Сердцебиение.docx',
    15: 'Фолликулометрия.docx',
}


def _set_run(para, run_idx: int, text: str):
    """Set text on a specific run if it exists."""
    if run_idx < len(para.runs):
        para.runs[run_idx].text = text


def _set_recom(para, value: str):
    """Recomendatsiya satri — agar bo'sh bo'lsa, butun paragrafni tozalaydi."""
    val = (value or '').strip()
    if not val:
        for r in para.runs:
            r.text = ''
    else:
        if para.runs:
            para.runs[0].text = f'Рекомендации: {val}'
            for r in para.runs[1:]:
                r.text = ''


def _fill_header(doc: Document, client: Dict, doctor: Dict, created_at: str):
    """Fill patient info, date, and doctor name paragraphs (common to all templates)."""
    last = client.get('last_name', '') or ''
    first = client.get('first_name', '') or ''
    patronymic = client.get('patronymic', '') or ''
    full_name = ' '.join(p for p in [last, first, patronymic] if p)

    birth_date = client.get('birth_date', '') or ''
    birth_year = str(birth_date)[:4] if birth_date else ''

    gender_val = client.get('gender', '') or ''
    gender_str = 'Мужской' if gender_val == 'male' else ('Женский' if gender_val == 'female' else gender_val)

    date_str = created_at[:10] if created_at else ''

    d_last = doctor.get('last_name', '') or ''
    d_first = doctor.get('first_name', '') or ''
    d_pat = doctor.get('patronymic_name', doctor.get('patronymic', '')) or ''
    doctor_name = ' '.join(p for p in [d_last, d_first, d_pat] if p)

    for p in doc.paragraphs:
        text = p.text
        if 'Ф.И.О. пациента' in text:
            _set_run(p, 0, f'Ф.И.О. пациента {full_name}')
            if 'Пол' in text:
                _set_run(p, 1, '.')
                _set_run(p, 2, f' Пол {gender_str}. {birth_year} г.р.')
            else:
                _set_run(p, 1, ',')
                _set_run(p, 2, f' {birth_year} г.р.')
        elif 'Дата проведения исследования:' in text:
            _set_run(p, 0, f'Дата проведения исследования: {date_str}')
        elif 'Дата обследования:' in text:
            _set_run(p, 0, f' Дата обследования: {date_str}')
        elif 'Врач_________________' in text:
            # Birinchi run ga matnni yozib, qolganlarini tozalaymiz (takrorlanishning oldini olish uchun)
            if p.runs:
                p.runs[0].text = f'Врач_________________ {doctor_name}.'
                for r in p.runs[1:]:
                    r.text = ''
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(40)
        elif 'Тел.:' in text or 'тел.:' in text.lower():
            # Leading bo'sh joylarni olib tashlab, faqat matn qoldiramiz
            if p.runs:
                for r in p.runs:
                    r.text = r.text.lstrip()
                    if r.text:
                        break
                # Qolgan runlarni tozalamaymiz, lekin trim qilamiz
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ─── Protocol 1: Печень и желчный пузырь ──────────────────────────────────────

def _fill_protocol_1(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # ── Печень ───────────────────────────────────────────────────────────────
    _set_run(paras[10], 2, fd.get('pech_topografia', ''))
    _set_run(paras[11], 1, f' {fd.get("pech_konturi", "")} ')
    _set_run(paras[12], 1, f' ПЗР левой доли {fd.get("pech_pzr_lev", "")} ')
    _set_run(paras[12], 7, f'ККР левой доли {fd.get("pech_kkr_lev", "")} ')
    _set_run(paras[13], 0, f'ПЗР хвостатой доли {fd.get("pech_pzr_hvost", "")} ')
    _set_run(paras[14], 0, f'ПЗР правой доли {fd.get("pech_pzr_prav", "")} ')
    _set_run(paras[14], 5, f'  ККР правой доли {fd.get("pech_kkr_prav", "")} ')
    _set_run(paras[15], 0, f'КВР {fd.get("pech_kvr", "")} мм ')
    _set_run(paras[15], 5, f'-латеральный размер {fd.get("pech_lat_razmer", "")} ')
    _set_run(paras[16], 1, f'л правой доли {fd.get("pech_ugol_prav", "")} ')
    _set_run(paras[16], 7, f' {fd.get("pech_ugol_lev", "")}')
    _set_run(paras[17], 1, f' {fd.get("pech_obem", "")} ')
    _set_run(paras[18], 0, f'Масса печени {fd.get("pech_massa", "")} ')
    _set_run(paras[19], 2, fd.get('pech_exostruktura', ''))
    _set_run(paras[20], 2, fd.get('pech_exogennost', ''))
    _set_run(paras[21], 0, f'Звукопроводимость {fd.get("pech_zvukoprov", "")} ')

    # ── Сосуды ───────────────────────────────────────────────────────────────
    _set_run(paras[23], 1, f' диаметром до {fd.get("soud_pech_veny", "")} мм ')
    _set_run(paras[24], 0, f'Суммарный диаметр {fd.get("soud_sum_diam", "")} мм')
    _set_run(paras[25], 0, f'Портальная вена {fd.get("soud_port_vena", "")} ')
    _set_run(paras[25], 4, f' Нижняя полая вена {fd.get("soud_nizhnyaya_pv", "")} ')

    # ── Желчные протоки ──────────────────────────────────────────────────────
    _set_run(paras[27], 1, f' {fd.get("zhp_vnutrip", "")} ')
    _set_run(paras[28], 0, f'Общий желчный {fd.get("zhp_obshchiy", "")} ')

    # ── Желчный пузырь ───────────────────────────────────────────────────────
    _set_run(paras[30], 2, f' {fd.get("puz_topografia", "")} ')
    _set_run(paras[31], 1, f' {fd.get("puz_forma", "")} ')
    _set_run(paras[32], 1, f' {fd.get("puz_kontur", "")} ')
    _set_run(paras[34], 0, f'Длина {fd.get("puz_dlina", "")} ')
    _set_run(paras[34], 4, f' Толщина {fd.get("puz_tolshina", "")} ')
    _set_run(paras[34], 8, f'). Ширина {fd.get("puz_shirina", "")} ')
    _set_run(paras[35], 1, f' {fd.get("puz_obem", "")} ')
    _set_run(paras[35], 9, f' {fd.get("puz_ploshad", "")} ')
    _set_run(paras[36], 0, f'Пузырный проток {fd.get("puz_protok", "")} ')
    _set_run(paras[37], 0, f'Толщина стенки {fd.get("puz_stenka", "")} ')
    _set_run(paras[38], 0, f'Содержимое {fd.get("puz_soderzhimoe", "")} ')

    # ── Заключение / Рекомендации ─────────────────────────────────────────────
    _set_run(paras[39], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_recom(paras[40], fd.get('rekomendacii', ''))


# ─── Protocol 2: Почки, мочеточники и мочевой пузырь ─────────────────────────
#
# Template structure (0-based):
#   [9]='Почки', [10]='Правая'
#   [11]=Топография, [12]=Подвижность, [13]=Жировая капсула, [14]=Фиброзная капсула
#   [15]=Контуры, [16]='Линейные размеры' HEADER
#   [17]=Длина/Толщина/Ширина (all in one para), [18]=Объём
#   [19]=Толщина паренхимы, [20]=Эхогенность, [21]=Эхоструктура
#   [22]=Состояние ЧЛС, [23]=КРИ, [24]=Конкременты, [25]=Мочеточник
#   [26]='', [27]='Левая'
#   [28]=Топография, [29]=Подвижность, [30]=Жировая капсула, [31]=Фиброзная капсула
#   [32]=Контуры, [33]='Линейные размеры' HEADER
#   [34]=Длина/Толщина/Ширина, [35]=Объём
#   [36]=Толщина паренхимы, [37]=Эхогенность, [38]=Эхоструктура
#   [39]=Состояние ЧЛС, [40]=КРИ, [41]=Конкременты, [42]=Мочеточник
#   [44]='Мочевой пузырь', [45]=Объём, [46]=Форма, [47]='Контур' HEADER
#   [48]=Наружный, [49]=Внутренний, [50]=Стенка, [51]=Содержимое
#   [52]=Шейка, [53]=Устья, [54]=Остаточная, [56]=Заключение

def _fill_protocol_2(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # ── Правая почка ──────────────────────────────────────────────────────────
    # [11] Топография: r[2]='в норме'
    _set_run(paras[11], 2, f' {fd.get("prav_topografia", "")}')
    # [12] Подвижность: r[1]=' сохранена  '
    _set_run(paras[12], 1, f' {fd.get("prav_podvizhnost", "")}  ')
    # [13] Жировая капсула: r[2]='без особенностей'
    _set_run(paras[13], 2, fd.get('prav_zhir_kapsul', ''))
    # [14] Фиброзная капсула: r[2]='без особенностей'
    _set_run(paras[14], 2, fd.get('prav_fib_kapsul', ''))
    # [15] Контуры: r[2]='в норме '
    _set_run(paras[15], 2, f'{fd.get("prav_konturi", "")} ')
    # [16] 'Линейные размеры' HEADER — skip
    # [17] Длина/Толщина/Ширина: r[1]=dlina blank, r[8]=combined tolshina, r[10]=combined shirina
    _set_run(paras[17], 1, f' {fd.get("prav_dlina", "")} ')
    _set_run(paras[17], 8, f'     Толщина {fd.get("prav_tolshina", "")} мм ')
    _set_run(paras[17], 10, f'   Ширина {fd.get("prav_shirina", "")} мм ')
    # [18] Объём: r[0]='Объём  '
    _set_run(paras[18], 0, f'Объём {fd.get("prav_obem", "")} ')
    # [19] Толщина паренхимы: r[1]='  '
    _set_run(paras[19], 1, f' {fd.get("prav_parenhima_t", "")} ')
    # [20] Эхогенность паренхимы: r[2]='сохранена'
    _set_run(paras[20], 2, fd.get('prav_exogen_par', ''))
    # [21] Эхоструктура паренхимы: r[2]='без особенностей.'
    _set_run(paras[21], 2, fd.get('prav_exostr_par', ''))
    # [22] Состояние ЧЛС: r[3]='толщина  ', r[7]=': без особенностей '
    _set_run(paras[22], 3, f'толщина {fd.get("prav_chls_tolshina", "")} ')
    _set_run(paras[22], 7, f': {fd.get("prav_chls_exostr", "")} ')
    # [23] КРИ: r[2]='   '
    _set_run(paras[23], 2, f' {fd.get("prav_kri", "")}')
    # [24] Конкременты: r[1]=': не визуализируются'
    _set_run(paras[24], 1, f': {fd.get("prav_konkrement", "")}')
    # [25] Мочеточник: r[2]='без особенностей'
    _set_run(paras[25], 2, fd.get('prav_mochetchnik', ''))

    # ── Левая почка ───────────────────────────────────────────────────────────
    # [28] Топография: r[2]='в норме '
    _set_run(paras[28], 2, f'{fd.get("lev_topografia", "")} ')
    # [29] Подвижность: r[0]='Дыхательная подвижность сохранена  '
    _set_run(paras[29], 0, f'Дыхательная подвижность {fd.get("lev_podvizhnost", "")}  ')
    # [30] Жировая капсула: r[0]='Жировая капсула без особенностей  '
    _set_run(paras[30], 0, f'Жировая капсула {fd.get("lev_zhir_kapsul", "")}  ')
    # [31] Фиброзная капсула: r[2]='без особенностей  '
    _set_run(paras[31], 2, f'{fd.get("lev_fib_kapsul", "")}  ')
    # [32] Контуры: r[0]='Контуры в норме '
    _set_run(paras[32], 0, f'Контуры {fd.get("lev_konturi", "")} ')
    # [33] 'Линейные размеры' HEADER — skip
    # [34] Длина/Толщина/Ширина: r[0]=dlina, r[5]=tolshina combined, r[7]=shirina combined
    _set_run(paras[34], 0, f'Длина {fd.get("lev_dlina", "")} ')
    _set_run(paras[34], 5, f')     Толщина {fd.get("lev_tolshina", "")} мм ')
    _set_run(paras[34], 7, f'   Ширина {fd.get("lev_shirina", "")} мм ')
    # [35] Объём: r[0]='Объём   '
    _set_run(paras[35], 0, f'Объём {fd.get("lev_obem", "")} ')
    # [36] Толщина паренхимы: r[0]='Толщина паренхимы  '
    _set_run(paras[36], 0, f'Толщина паренхимы {fd.get("lev_parenhima_t", "")} ')
    # [37] Эхогенность паренхимы: r[1]='сохранена'
    _set_run(paras[37], 1, fd.get('lev_exogen_par', ''))
    # [38] Эхоструктура паренхимы: r[0]='Эхоструктура паренхимы без особенностей.'
    _set_run(paras[38], 0, f'Эхоструктура паренхимы {fd.get("lev_exostr_par", "")}.')
    # [39] Состояние ЧЛС: r[0]='Состояние ЧЛС: толщина  ', r[4]=': без особенностей '
    _set_run(paras[39], 0, f'Состояние ЧЛС: толщина {fd.get("lev_chls_tolshina", "")} ')
    _set_run(paras[39], 4, f': {fd.get("lev_chls_exostr", "")} ')
    # [40] КРИ: r[1]='-ренальный индекс   '
    _set_run(paras[40], 1, f'-ренальный индекс {fd.get("lev_kri", "")} ')
    # [41] Конкременты: r[0]='Наличие конкрементов в ЧЛС: не визуализи...'
    _set_run(paras[41], 0, f'Наличие конкрементов в ЧЛС: {fd.get("lev_konkrement", "")}')
    # [42] Мочеточник: r[1]=' без особенностей '
    _set_run(paras[42], 1, f' {fd.get("lev_mochetchnik", "")} ')

    # ── Мочевой пузырь ────────────────────────────────────────────────────────
    # [45] Объём: r[0]='Объём  мл '
    _set_run(paras[45], 0, f'Объём {fd.get("mp_obem", "")} мл ')
    # [46] Форма: r[0]='Форма обычная '
    _set_run(paras[46], 0, f'Форма {fd.get("mp_forma", "")} ')
    # [47] 'Контур' HEADER — skip
    # [48] Наружный контур: r[2]=' в норме '
    _set_run(paras[48], 2, f' {fd.get("mp_kontur_naruzh", "")} ')
    # [49] Внутренний контур: r[2]=' в норме '
    _set_run(paras[49], 2, f' {fd.get("mp_kontur_vnutr", "")} ')
    # [50] Толщина стенки: r[0]='...до опорожнения  ', r[4]='),...после опорожнения  мм '
    _set_run(paras[50], 0, f'Толщина стенки: до опорожнения {fd.get("mp_stenka_do", "")} ')
    _set_run(paras[50], 4, f'), после опорожнения {fd.get("mp_stenka_posle", "")} мм ')
    # [51] Содержимое: r[0]='Содержимое анэхогенное '
    _set_run(paras[51], 0, f'Содержимое {fd.get("mp_soderzhimoe", "")} ')
    # [52] Шейка: r[0]='Состояние шейки, зева в норме '
    _set_run(paras[52], 0, f'Состояние шейки, зева {fd.get("mp_sheika", "")} ')
    # [53] Устья: r[0]='Локализация устьев мочеточников в норме '
    _set_run(paras[53], 0, f'Локализация устьев мочеточников {fd.get("mp_ustya", "")} ')
    # [54] Остаточная: r[1]='  мл '
    _set_run(paras[54], 1, f' {fd.get("mp_ostatoch", "")} мл ')

    # ── Заключение ────────────────────────────────────────────────────────────
    _set_run(paras[56], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')


# ─── Protocol 3: Щитовидная железа ───────────────────────────────────────────

def _fill_protocol_3(doc: Document, fd: Dict):
    paras = doc.paragraphs

    _set_run(paras[10], 2, fd.get('p_topografiya', ''))
    _set_run(paras[11], 2, fd.get('p_kontur', ''))
    _set_run(paras[12], 2, fd.get('p_elastichnost', ''))
    _set_run(paras[13], 2, fd.get('p_podvijnost', ''))
    _set_run(paras[14], 2, f' {fd.get("p_razmer", "")} ')
    _set_run(paras[15], 2, fd.get('p_ehostruktura', ''))
    _set_run(paras[16], 2, fd.get('p_ehogennost', ''))

    _set_run(paras[19], 2, f'{fd.get("dp_topografiya", "")} ')
    _set_run(paras[20], 2, fd.get('dp_kontur', ''))
    _set_run(paras[21], 1, f'  {fd.get("dp_elastichnost", "")}')
    _set_run(paras[22], 1, fd.get('dp_podvijnost', ''))
    _set_run(paras[23], 3, f' длина {fd.get("dp_dlina", "")} ')
    _set_run(paras[23], 7, f'лщина {fd.get("dp_tolshina", "")} мм,')
    _set_run(paras[23], 8, f' ширина {fd.get("dp_shirina", "")} мм.   Объём {fd.get("dp_obem", "")} ')
    _set_run(paras[24], 2, f'{fd.get("dp_forma", "")} ')
    _set_run(paras[25], 3, fd.get('dp_ehostruktura', ''))
    _set_run(paras[26], 2, fd.get('dp_ehogennost', ''))

    _set_run(paras[29], 2, fd.get('dl_topografiya', ''))
    _set_run(paras[30], 2, fd.get('dl_kontur', ''))
    _set_run(paras[31], 2, fd.get('dl_elastichnost', ''))
    _set_run(paras[32], 2, f' {fd.get("dl_podvijnost", "")}')
    _set_run(paras[33], 1, f': длина {fd.get("dl_dlina", "")} мм, т')
    _set_run(paras[33], 3, f'лщина {fd.get("dl_tolshina", "")} мм')
    _set_run(paras[33], 4, f', ширина {fd.get("dl_shirina", "")} мм.   Объём {fd.get("dl_obem", "")}')
    _set_run(paras[34], 1, f' {fd.get("dl_forma", "")}')
    _set_run(paras[35], 2, fd.get('dl_ehostruktura', ''))
    _set_run(paras[36], 2, fd.get('dl_ehogennost', ''))

    _set_run(paras[38], 0, f'Суммарный объём {fd.get("summarniy", "")} ')
    _set_run(paras[40], 0, f'Объём железы  в норме у пациента при весе {fd.get("kg", "")} ')
    _set_run(paras[40], 2, f' составляет {fd.get("sm", "")} куб. см.')

    _set_run(paras[41], 0, f'Заключение: {fd.get("zaklyuchenie_sh", "")}')
    _set_run(paras[41], 1, '')
    _set_recom(paras[42], fd.get('rekomendatsi_sh', ''))
    _set_run(paras[42], 1, '')


# ─── Protocol 4: Поджелудочная железа ────────────────────────────────────────

def _fill_protocol_4(doc: Document, fd: Dict):
    paras = doc.paragraphs

    _set_run(paras[9], 0, f'Топография {fd.get("pdj_topografia", "")}')
    _set_run(paras[10], 0, f'Форма {fd.get("pdj_forma", "")}')
    _set_run(paras[11], 0, f'Контур {fd.get("pdj_kontur", "")}')
    _set_run(paras[13], 0, f'Толщина головки {fd.get("pdj_golovka", "")} ')
    _set_run(paras[13], 4, f' тела {fd.get("pdj_telo", "")} мм ')
    _set_run(paras[13], 7, f', хвоста {fd.get("pdj_hvost", "")} ')
    _set_run(paras[14], 0, f'Длина {fd.get("pdj_dlina", "")} ')
    _set_run(paras[14], 1, '')
    _set_run(paras[15], 2, fd.get('pdj_exostruktura', ''))
    _set_run(paras[16], 0, f'Эхогенность {fd.get("pdj_exogennost", "")} ')
    _set_run(paras[17], 1, f'{fd.get("pdj_zvukoprov", "")} ')
    _set_run(paras[18], 2, f'сти головки {fd.get("pdj_virsungov_g", "")} ')
    _set_run(paras[18], 3, '')
    _set_run(paras[19], 1, f'ти тела {fd.get("pdj_virsungov_t", "")} ')
    _set_run(paras[20], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_recom(paras[21], fd.get('rekomendacii', ''))


# ─── Protocol 5: Селезёнка ────────────────────────────────────────────────────

def _fill_protocol_5(doc: Document, fd: Dict):
    paras = doc.paragraphs

    _set_run(paras[8], 1, f' {fd.get("sel_topografia", "")}')
    _set_run(paras[9], 1, f'\t {fd.get("sel_forma", "")} ')
    _set_run(paras[10], 0, f'Контуры {fd.get("sel_konturi", "")} ')
    _set_run(paras[11], 0, f'Размеры линейные.  Длина {fd.get("sel_dlina", "")} ')
    _set_run(paras[12], 1, f'                          Толщина {fd.get("sel_tolshina", "")} ')
    _set_run(paras[13], 1, f'                          Ширина {fd.get("sel_shirina", "")} ')
    _set_run(paras[14], 1, f'                                Объём {fd.get("sel_obem", "")} см')
    _set_run(paras[15], 0, f'                                    Селезёночный индекс {fd.get("sel_indeks", "")} см')
    _set_run(paras[16], 0, f'Эхоструктура {fd.get("sel_exostruktura", "")}')
    _set_run(paras[17], 0, f'Эхогенность {fd.get("sel_exogennost", "")} ')
    _set_run(paras[18], 2, f'{fd.get("sel_zvukoprov", "")} ')
    _set_run(paras[19], 3, f' {fd.get("sel_lienalis", "")} ')
    _set_run(paras[21], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_recom(paras[22], fd.get('rekomendacii', ''))


# ─── Protocol 6: Простата и семенные пузырьки ────────────────────────────────
#
# Template:
#   [10]=Топография, [11]=Форма (r[0]='Ф', r[1]='орма в норме'), [12]=Контур
#   [13]=Капсула толщина (r[1]=blank), [14]=Размеры (r[1]=dlina, r[6]=tolshina, r[9]=shirina)
#   [15]=Объём (r[0]='Объем  см'), [16]=Индекс (r[2]=blank), [17]=Эхоструктура (r[1])
#   [18]=Эхогенность (r[0]), [19]=Уретра (r[1]='канала в норме')
#   [21]=МП (r[1]=combined)
#   [26]=Топография пузырька, [27]=Контур (r[0]='Контур', r[1]=' в норме')
#   [28]=Размеры (r[0]=dlina, r[4]=tolshina, r[8]=shirina suffix 'рина')
#   [29]=Эхоструктура (r[2]='о эякуляции  ', r[7]='после эякуляции')
#   [30]=Эхогенность (r[0])
#   [32]=Заключение, [33]=Рекомендации

def _fill_protocol_6(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # ── Предстательная железа ─────────────────────────────────────────────────
    _set_run(paras[10], 0, f'Топография {fd.get("pr_topografia", "")}')
    _set_run(paras[11], 1, f'орма {fd.get("pr_forma", "")}')
    _set_run(paras[12], 0, f'Контур {fd.get("pr_kontur", "")}')
    _set_run(paras[13], 1, f' {fd.get("pr_kapsul_t", "")} ')
    _set_run(paras[14], 1, f'е: длина {fd.get("pr_dlina", "")} ')
    _set_run(paras[14], 6, f',толщина {fd.get("pr_tolshina", "")} мм (в норме ')
    _set_run(paras[14], 9, f'ширина {fd.get("pr_shirina", "")} мм ')
    _set_run(paras[15], 0, f'Объем {fd.get("pr_obem", "")} см')
    _set_run(paras[16], 2, f' {fd.get("pr_indeks", "")} ')
    _set_run(paras[17], 1, f' {fd.get("pr_exostruktura", "")} ')
    _set_run(paras[18], 0, f'Эхогенность {fd.get("pr_exogennost", "")} ')
    _set_run(paras[19], 1, f'канала {fd.get("pr_urethra", "")}')
    _set_run(paras[21], 1,
             f'й объем мочевого пузыря {fd.get("pr_nach_obem", "")} мл.  '
             f'Остаточная моча {fd.get("pr_ostatoch", "")} мл')

    # ── Семенные пузырьки ─────────────────────────────────────────────────────
    _set_run(paras[26], 0, f'Топография {fd.get("svp_topografia", "")}')
    # [27] Контур: r[0]='Контур', r[1]=' в норме' → replace r[1] to avoid duplicate
    _set_run(paras[27], 1, f' {fd.get("svp_kontur", "")}')
    # [28] Размеры: r[0]='Размеры: длина   ', r[4]='толщина  мм ', r[7]=' ши', r[8]='рина   мм '
    _set_run(paras[28], 0, f'Размеры: длина {fd.get("svp_dlina", "")} ')
    _set_run(paras[28], 4, f'толщина {fd.get("svp_tolshina", "")} мм ')
    _set_run(paras[28], 8, f'рина {fd.get("svp_shirina", "")} мм ')
    # [29] Эхоструктура: r[2]='о эякуляции   '(blank), r[7]='после эякуляции'
    _set_run(paras[29], 2, f'о эякуляции {fd.get("svp_exostr_do", "")} ')
    _set_run(paras[29], 7, f'после {fd.get("svp_exostr_posle", "")}')
    _set_run(paras[30], 0, f'Эхогенность {fd.get("svp_exogennost", "")}')

    # Левый семенной пузырёк — paragraf 30 dan keyin yangi paragraflar qo'shiladi
    has_lev = any(fd.get(k, '').strip() for k in ['svl_topografia', 'svl_kontur', 'svl_dlina', 'svl_tolshina', 'svl_shirina', 'svl_exostr_do', 'svl_exostr_posle', 'svl_exogennost'])
    if has_lev:
        from copy import deepcopy
        from docx.text.paragraph import Paragraph
        # Sarlavha paragrafi (Левый)
        title_el = deepcopy(paras[25]._element)
        paras[30]._element.addnext(title_el)
        title_p = Paragraph(title_el, paras[30]._parent)
        for r in title_p.runs:
            r.text = ''
        if title_p.runs:
            title_p.runs[0].text = '            Левый'
        # Qator-qator ma'lumotlar
        lines = [
            f'Топография {fd.get("svl_topografia", "")} (в норме латеральнее над простатой).',
            f'Контур {fd.get("svl_kontur", "")} (в норме дольчатый, четкий).',
            f'Размеры: длина {fd.get("svl_dlina", "")} мм (в норме до 50), толщина {fd.get("svl_tolshina", "")} мм (в норме до 15), ширина {fd.get("svl_shirina", "")} мм (в норме до 20).',
            f'Эхоструктура до эякуляции {fd.get("svl_exostr_do", "")} (в норме однородная, ячеистая), после эякуляции {fd.get("svl_exostr_posle", "")} (в норме однородная, гомогенная).',
            f'Эхогенность {fd.get("svl_exogennost", "")} (в норме гипоэхогенная).',
        ]
        prev_el = title_el
        for line in lines:
            new_el = deepcopy(paras[26]._element)
            prev_el.addnext(new_el)
            new_p = Paragraph(new_el, paras[30]._parent)
            for r in new_p.runs:
                r.text = ''
            if new_p.runs:
                new_p.runs[0].text = line
            prev_el = new_el

    _set_run(paras[32], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_recom(paras[33], fd.get('rekomendacii', ''))


# ─── Protocol 7: Молочные железы ─────────────────────────────────────────────
#
# Template (RIGHT [8-24], LEFT [26-42]):
#   [8]='ПРАВАЯ', [9]=Топография (r[0]='Топография ', r[1]='в норме')
#   [10]=Кожа контур (r[0]='Кожа: контур в норме')
#   [11]=Кожа толщина (r[0]='           Толщина ', r[1]=blank, r[2]='мм')
#   [12]=Кожа эхогенность (r[1]='          Эхогенность в норме')
#   [13]=Кожа эхоструктура (r[1]='       Эхоструктура в норме')
#   [14]=Сосок (r[0]='Сосок: без особенностей.')
#   [15]='Жировая ткань.' HEADER
#   [16]=Подкожная толщина (r[0]='          Подкожная, толщина  ')
#   [17]=Интрамаммарная (template label only, not editable here)
#   [18]=Ретромаммарная толщина (r[2]=', толщиной  ')
#   [19]=Паренхима контур (r[3]='  контур в норме')
#   [20]=Паренхима толщина (r[1]='       Толщина  ')
#   [21]=Паренхима эхогенность (r[1]='      Эхогенность в норме ')
#   [22]=Паренхима эхоструктура (r[1]=' без особенностей ')
#   [23]=Галактофоры (r[2]=' не визуализируются.')
#   [24]=Лимфузлы (r[1]='без особенностей.')
#   [25]='', [26]='ЛЕВАЯ'
#   LEFT side run structure differs:
#   [27]=Топография (r[0]='Топография', r[1]=' в ', r[2]='норме...') → use r[1], clear r[2-4]
#   [28]=Кожа контур (r[0]='Кожа: контур в норме ')
#   [29]=Кожа толщина (r[0]='           Толщина  ')
#   [30]=Кожа эхогенность (r[0]='           Эхогенность в норме ')
#   [31]=Кожа эхоструктура (r[0]='           Эхоструктура в норме ')
#   [32]=Сосок (r[0]='Сосок: без особенностей.')
#   [33]='Жировая ткань.' HEADER
#   [34]=Подкожная толщина (r[0]='          Подкожная, толщина  ')
#   [35]=Интрамаммарная (label only)
#   [36]=Ретромаммарная (r[2]=', толщиной  ')
#   [37]=Паренхима контур (r[1]=' Железистая ткань:  контур в норме ')
#   [38]=Паренхима толщина (r[0]='             Толщина  ')
#   [39]=Паренхима эхогенность (r[0]='             Эхогенность в норме ')
#   [40]=Паренхима эхоструктура (r[0]='             Эхоструктура без...')
#   [41]=Галактофоры (r[2]=' не визуализируются.')
#   [42]=Лимфузлы (r[1]=' без особенностей.')
#   [44]=Заключение, [45]=Рекомендации

def _fill_protocol_7(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # ── Правая молочная железа ────────────────────────────────────────────────
    _set_run(paras[9], 1, fd.get('prm_topografia', ''))
    _set_run(paras[10], 0, f'Кожа: контур {fd.get("prm_kozha_kontur", "")}')
    # [11] kozha_t: r[1]=blank (between 'Толщина ' and 'мм')
    _set_run(paras[11], 1, f'{fd.get("prm_kozha_t", "")} ')
    _set_run(paras[12], 1, f'          Эхогенность {fd.get("prm_kozha_exogen", "")}')
    _set_run(paras[13], 1, f'       Эхоструктура {fd.get("prm_kozha_exostr", "")}')
    _set_run(paras[14], 0, f'Сосок: {fd.get("prm_sosok", "")}.')
    # [15]='Жировая ткань.' HEADER — skip
    _set_run(paras[16], 0, f'          Подкожная, толщина {fd.get("prm_podkozh_t", "")} ')
    _set_run(paras[16], 2, '')  # clear ", без особенностей."
    # [17]=Интрамаммарная HEADER — skip
    _set_run(paras[18], 2, f', толщиной {fd.get("prm_retromam_t", "")} ')
    _set_run(paras[18], 4, '')  # clear ", без особенностей."
    _set_run(paras[19], 3, f'  контур {fd.get("prm_parenhima_kontur", "")}')
    _set_run(paras[20], 1, f'       Толщина {fd.get("prm_parenhima_t", "")} ')
    _set_run(paras[21], 1, f'      Эхогенность {fd.get("prm_parenhima_exogen", "")} ')
    _set_run(paras[22], 1, f' {fd.get("prm_parenhima_exostr", "")} ')
    _set_run(paras[23], 2, f' {fd.get("prm_galaktofori", "")}.')
    # Особенности правой железы — Лимфатические узлы dan oldin
    prm_osob = (fd.get('prm_osobennosti', '') or '').strip()
    if prm_osob:
        from copy import deepcopy
        from docx.text.paragraph import Paragraph
        new_el = deepcopy(paras[24]._element)
        paras[24]._element.addprevious(new_el)
        osb_p = Paragraph(new_el, paras[24]._parent)
        for r in osb_p.runs:
            r.text = ''
        if osb_p.runs:
            osb_p.runs[0].text = f'Особенности: {prm_osob}'
    _set_run(paras[24], 1, f'{fd.get("prm_limfuzli", "")}.')

    # ── Левая молочная железа ─────────────────────────────────────────────────
    # [27] Топография: r[1]=' в ', r[2]='норме...' → put value in r[1], clear r[2-4]
    _set_run(paras[27], 1, f' {fd.get("levm_topografia", "")}')
    _set_run(paras[27], 2, '')
    _set_run(paras[27], 3, '')
    _set_run(paras[27], 4, '')
    # [28] Кожа контур: r[0]='Кожа: контур в норме '
    _set_run(paras[28], 0, f'Кожа: контур {fd.get("levm_kozha_kontur", "")} ')
    # [29] Кожа толщина: r[0]='           Толщина  '
    _set_run(paras[29], 0, f'           Толщина {fd.get("levm_kozha_t", "")} ')
    # [30] Кожа эхогенность: r[0]='           Эхогенность в норме '
    _set_run(paras[30], 0, f'           Эхогенность {fd.get("levm_kozha_exogen", "")} ')
    # [31] Кожа эхоструктура: r[0]='           Эхоструктура в норме '
    _set_run(paras[31], 0, f'           Эхоструктура {fd.get("levm_kozha_exostr", "")} ')
    _set_run(paras[32], 0, f'Сосок: {fd.get("levm_sosok", "")}.')
    # [33]='Жировая ткань.' HEADER — skip
    _set_run(paras[34], 0, f'          Подкожная, толщина {fd.get("levm_podkozh_t", "")} ')
    _set_run(paras[34], 2, '')  # clear ", без особенностей."
    # [35]=Интрамаммарная HEADER — skip
    _set_run(paras[36], 2, f', толщиной {fd.get("levm_retromam_t", "")} ')
    _set_run(paras[36], 4, '')  # clear ", без особенностей."
    _set_run(paras[37], 1, f' Железистая ткань:  контур {fd.get("levm_parenhima_kontur", "")} ')
    _set_run(paras[38], 0, f'             Толщина {fd.get("levm_parenhima_t", "")} ')
    _set_run(paras[39], 0, f'             Эхогенность {fd.get("levm_parenhima_exogen", "")} ')
    _set_run(paras[40], 0, f'             Эхоструктура {fd.get("levm_parenhima_exostr", "")} ')
    _set_run(paras[41], 2, f' {fd.get("levm_galaktofori", "")}.')
    # Особенности левой железы — Лимфатические узлы dan oldin
    levm_osob = (fd.get('levm_osobennosti', '') or '').strip()
    if levm_osob:
        from copy import deepcopy
        from docx.text.paragraph import Paragraph
        new_el = deepcopy(paras[42]._element)
        paras[42]._element.addprevious(new_el)
        osb_p = Paragraph(new_el, paras[42]._parent)
        for r in osb_p.runs:
            r.text = ''
        if osb_p.runs:
            osb_p.runs[0].text = f'Особенности: {levm_osob}'
    _set_run(paras[42], 1, f' {fd.get("levm_limfuzli", "")}.')

    _set_run(paras[44], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_recom(paras[45], fd.get('rekomendacii', ''))


# ─── Protocol 8: Мочевой пузырь ──────────────────────────────────────────────

def _fill_protocol_8(doc: Document, fd: Dict):
    paras = doc.paragraphs

    _set_run(paras[8], 1, f'бъем, физиологический {fd.get("mp_obem_fiz", "")} ')
    _set_run(paras[9], 0, f'Форма {fd.get("mp_forma", "")} ')
    _set_run(paras[10], 3, fd.get('mp_kontur_naruzh', ''))
    _set_run(paras[11], 3, fd.get('mp_kontur_vnutr', ''))
    _set_run(paras[12], 1, f'енки:   до опорожнения {fd.get("mp_stenka_do", "")} ')
    _set_run(paras[12], 7, f'  после опорожнения {fd.get("mp_stenka_posle", "")} мм ')
    _set_run(paras[13], 1, f'одержимое {fd.get("mp_soderzhimoe", "")} ')
    _set_run(paras[14], 1, f' {fd.get("mp_sheika", "")} ')
    _set_run(paras[15], 0, f'Локализация устьев мочеточников {fd.get("mp_ustya", "")} ')
    _set_run(paras[16], 1, f'аточной мочи {fd.get("mp_ostatoch", "")} ')
    _set_run(paras[16], 2, ' мл ')
    _set_run(paras[18], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_recom(paras[19], fd.get('rekomendacii', ''))


# ─── Protocol 9: Органы малого таза ──────────────────────────────────────────
#
# Key run corrections vs initial code:
#   [9] den_cikla: r[3] (not r[4])
#   [21] sheika_tolshina: r[7]=f' {value} мм ' (not just value)
#   [42] yach_lev_kontur: r[0]=f'Контур {value} ' (not r[1])
#   [43] yach_lev_tolshina: r[4] (not r[5]); shirina: r[6] (not r[7])
#   [45] yach_lev_exostr: r[0] (not r[1])
#   [46] yach_lev_follikuli: r[0] (not r[1])

def _fill_protocol_9(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # ── Матка ─────────────────────────────────────────────────────────────────
    # [9] Дата менструации: r[1]='  ,', r[3]='день менструального цикла'
    _set_run(paras[9], 1, f' {fd.get("mat_data_mens", "")},')
    _set_run(paras[9], 3, f'{fd.get("mat_den_cikla", "")} день менструального цикла')
    # Менопауза — agar yozilgan bo'lsa, paragraph 9 dan keyin qo'shiladi
    menopauza = (fd.get('mat_menopauza', '') or '').strip()
    if menopauza:
        from copy import deepcopy
        from docx.text.paragraph import Paragraph
        new_el = deepcopy(paras[9]._element)
        paras[9]._element.addnext(new_el)
        mp_p = Paragraph(new_el, paras[9]._parent)
        for r in mp_p.runs:
            r.text = ''
        if mp_p.runs:
            mp_p.runs[0].text = f'Менопауза {menopauza} лет'
    # [10] Позиция/Положение/Форма
    _set_run(paras[10], 1, fd.get('mat_poziciya', ''))
    _set_run(paras[10], 5, fd.get('mat_polozhenie', ''))
    _set_run(paras[10], 8, f'{fd.get("mat_forma", "")} ')
    # [11] Соотношение тела матки к шейке
    _set_run(paras[11], 1, f' {fd.get("mat_sootn", "")}.')
    # [12] Размеры: r[1]=' мм ', r[6]='толщина  мм ', r[9]='; ширина  мм '
    _set_run(paras[12], 1, f' {fd.get("mat_dlina", "")} мм ')
    _set_run(paras[12], 6, f'толщина {fd.get("mat_tolshina", "")} мм ')
    _set_run(paras[12], 9, f'; ширина {fd.get("mat_shirina", "")} мм ')

    # толщина передней/задней стенки миометрия — [12] dan keyin yangi paragraf
    peredney = str(fd.get('mat_peredney', '') or '').strip()
    zadney = str(fd.get('mat_zadney', '') or '').strip()
    from copy import deepcopy
    from docx.text.paragraph import Paragraph
    stenki_el = deepcopy(paras[12]._element)
    paras[12]._element.addnext(stenki_el)
    stenki_p = Paragraph(stenki_el, paras[12]._parent)
    for r in stenki_p.runs:
        r.text = ''
    if stenki_p.runs:
        stenki_p.runs[0].text = f'Толщина передней стенки миометрия {peredney} мм, задней стенки миометрия {zadney} мм.'

    # [13] Объём: r[1]='          Объем матки   см'
    _set_run(paras[13], 1, f'          Объем матки {fd.get("mat_obem", "")} см')
    # [14] Контур: r[0]='Контур без изменений '
    _set_run(paras[14], 0, f'Контур {fd.get("mat_kontur", "")} ')
    # [15] Эхоструктура миометрия: r[3]='однородная'
    _set_run(paras[15], 3, fd.get('mat_mio_exostr', ''))
    # [16] Эхогенность миометрия: r[4]='средняя'
    _set_run(paras[16], 4, fd.get('mat_mio_exogen', ''))
    # [17] М-эхо: r[0]='Толщина эндометрия (М-эхо)  ', r[4]=' faza '
    _set_run(paras[17], 0, f'Толщина эндометрия (М-эхо) {fd.get("mat_endo_t", "")} ')
    _set_run(paras[17], 4, f' {fd.get("mat_endo_faza", "")} ')
    # [18] Эхоструктура эндометрия: r[2]='однородная'
    _set_run(paras[18], 2, fd.get('mat_endo_exostr', ''))
    # [19] Полость матки: r[0]='Полость матки в норме '
    _set_run(paras[19], 0, f'Полость матки {fd.get("mat_polost", "")} ')

    # ── Шейка матки ───────────────────────────────────────────────────────────
    # [21] Размеры: r[1]=dlina blank, r[7]=' мм '(tolshina+мм), r[11]=shirina blank
    _set_run(paras[21], 1, fd.get('sheika_dlina', ''))
    _set_run(paras[21], 7, f' {fd.get("sheika_tolshina", "")} мм ')
    _set_run(paras[21], 11, fd.get('sheika_shirina', ''))
    _set_run(paras[22], 1, f'нтур {fd.get("sheika_kontur", "")} ')
    _set_run(paras[23], 1, f'енки {fd.get("sheika_exostr", "")} ')
    _set_run(paras[24], 1, f'ь стенки {fd.get("sheika_exogen", "")}')
    _set_run(paras[25], 4, f' {fd.get("sheika_endo_t", "")} мм ')
    _set_run(paras[26], 3, f' {fd.get("sheika_endo_exostr", "")}.')
    _set_run(paras[27], 1, f' {fd.get("sheika_polost", "")}')

    # ── Правый яичник ─────────────────────────────────────────────────────────
    _set_run(paras[31], 0, f'Топография {fd.get("yach_pr_topografia", "")}')
    _set_run(paras[32], 0, f'Форма {fd.get("yach_pr_forma", "")} ')
    _set_run(paras[33], 1, f' {fd.get("yach_pr_kontur", "")} ')
    _set_run(paras[34], 0, f'Размеры:  длина {fd.get("yach_pr_dlina", "")} ')
    _set_run(paras[34], 5, f'ина {fd.get("yach_pr_tolshina", "")} мм ')
    _set_run(paras[34], 7, f'  ширина {fd.get("yach_pr_shirina", "")} мм ')
    _set_run(paras[35], 0, f'Объем {fd.get("yach_pr_obem", "")} см')
    _set_run(paras[36], 1, f' (фолликулярный аппарат) {fd.get("yach_pr_exostr", "")}.')
    _set_run(paras[37], 1, f' {fd.get("yach_pr_follikuli", "")} ')
    _set_run(paras[38], 0, f'Эхогенность {fd.get("yach_pr_exogen", "")} ')

    # ── Левый яичник ──────────────────────────────────────────────────────────
    _set_run(paras[40], 0, f'Топография {fd.get("yach_lev_topografia", "")}')
    _set_run(paras[41], 0, f'Форма {fd.get("yach_lev_forma", "")} ')
    # [42] Контур: r[0]='Контур в норме ' (single run for left side)
    _set_run(paras[42], 0, f'Контур {fd.get("yach_lev_kontur", "")} ')
    # [43] Размеры: r[0]=dlina, r[4]=tolshina combined, r[6]=shirina combined
    _set_run(paras[43], 0, f'Размеры:  длина {fd.get("yach_lev_dlina", "")} ')
    _set_run(paras[43], 4, f',  толщина {fd.get("yach_lev_tolshina", "")} мм ')
    _set_run(paras[43], 6, f'ширина {fd.get("yach_lev_shirina", "")} мм ')
    _set_run(paras[44], 0, f'Объем {fd.get("yach_lev_obem", "")} см')
    # [45] Эхоструктура: r[0] is single run for left side
    _set_run(paras[45], 0, f'Эхоструктура (фолликулярный аппарат) {fd.get("yach_lev_exostr", "")}. ')
    # [46] Фолликулы: r[0] is single run for left side
    _set_run(paras[46], 0, f'Фолликулы диаметром {fd.get("yach_lev_follikuli", "")} ')
    _set_run(paras[47], 0, f'Эхогенность {fd.get("yach_lev_exogen", "")} ')

    # ── Маточные трубы и дугласово ────────────────────────────────────────────
    _set_run(paras[50], 1, f' {fd.get("mat_truba_prav", "")} ')
    _set_run(paras[51], 1, f' {fd.get("mat_truba_lev", "")} ')
    _set_run(paras[52], 2, fd.get('duglas', ''))

    _set_run(paras[54], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_recom(paras[55], fd.get('rekomendacii', ''))


# ─── Protocol 10: Надпочечники ────────────────────────────────────────────────

def _fill_protocol_10(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # ── Правый надпочечник ────────────────────────────────────────────────────
    _set_run(paras[10], 0, f'Топография {fd.get("npr_topografia", "")}')
    _set_run(paras[11], 0, f'Форма {fd.get("npr_forma", "")}')
    _set_run(paras[13], 1, f'          Продольный {fd.get("npr_prodolniy", "")} ')
    _set_run(paras[14], 1, f'         Поперечный {fd.get("npr_poperechniy", "")} ')
    _set_run(paras[15], 1, f'           Толщина {fd.get("npr_tolshina", "")} ')
    _set_run(paras[16], 1, f'        Площадь {fd.get("npr_ploshad", "")} см')
    _set_run(paras[17], 0, f'          Индекс гиперплазии {fd.get("npr_giper_indeks", "")} ')
    _set_run(paras[18], 0, f'Эхоструктура {fd.get("npr_exostruktura", "")} ')
    _set_run(paras[19], 0, f'Эхогенность {fd.get("npr_exogennost", "")}')

    # ── Левый надпочечник ─────────────────────────────────────────────────────
    _set_run(paras[22], 0, f'Топография {fd.get("nlev_topografia", "")}')
    _set_run(paras[23], 0, f'Форма {fd.get("nlev_forma", "")}')
    _set_run(paras[25], 1, f'           Продольный {fd.get("nlev_prodolniy", "")} ')
    _set_run(paras[26], 1, f'           Поперечный {fd.get("nlev_poperechniy", "")} ')
    _set_run(paras[27], 1, f'            Толщина {fd.get("nlev_tolshina", "")} ')
    _set_run(paras[28], 1, f'         Площадь {fd.get("nlev_ploshad", "")} см')
    _set_run(paras[29], 0, f'          Индекс гиперплазии {fd.get("nlev_giper_indeks", "")} ')
    _set_run(paras[30], 0, f'Структура {fd.get("nlev_exostruktura", "")} ')
    _set_run(paras[31], 0, f'Эхогенность {fd.get("nlev_exogennost", "")}')

    _set_run(paras[33], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_recom(paras[34], fd.get('rekomendacii', ''))


# ─── Protocol 11: Коленный сустав ────────────────────────────────────────────
#
# Meniscus paras have side-by-side layout (передний+задний in one paragraph):
#   [29]=header, [30]=Форма (r[1]=per, r[7]=zad)
#   [31]=Контур (r[1]=f' Контур {per}', r[5]=f'...Контур {zad}')
#   [32]=Эхоструктура (r[1]=f'труктура {per} ', r[9]=f'Эхоструктура {zad} ')
#   [33]=Эхогенность (r[0]=f'  Эхогенность {per}', r[8]=f' Эхогенность {zad}')
#   [35]=header, [36]=Форма (r[1]=per, r[6]=zad)
#   [37]=Контур (r[0]=f'  Контур {per}', r[8]=f'Контур {zad}')
#   [38]=Эхоструктура (r[0]=f'  Эхоструктура {per} ', r[8]=f'Эхоструктура {zad} ')
#   [39]=Эхогенность (r[0]=f'  Эхогенность {per}', r[10]=f'  Эхогенность {zad} ')

def _fill_protocol_11(doc: Document, fd: Dict):
    paras = doc.paragraphs

    _set_run(paras[8], 0, f'Сустав {fd.get("kol_storona", "")}')

    # ── Собственная связка надколенника ──────────────────────────────────────
    _set_run(paras[10], 1, f'нтур {fd.get("kol_svyaz_kontur", "")}')
    _set_run(paras[11], 1, f' {fd.get("kol_svyaz_tolshina", "")} ')
    _set_run(paras[12], 1, f'руктура {fd.get("kol_svyaz_exostr", "")}')
    _set_run(paras[13], 1, f'ть {fd.get("kol_svyaz_exogen", "")}')

    # ── Связки и мягкие ткани ────────────────────────────────────────────────
    _set_run(paras[14], 0, f'Медиальная связка {fd.get("kol_med_svyaz", "")}.')
    _set_run(paras[15], 0, f'Латеральная связка {fd.get("kol_lat_svyaz", "")}.')
    _set_run(paras[16], 1, f'Гоффа {fd.get("kol_goffa", "")}.')
    _set_run(paras[17], 0, f'Мышцы, суставные сумки, связки  ')
    _set_run(paras[17], 1, f'{fd.get("kol_mishci", "")}.')

    # ── Суставная жидкость ────────────────────────────────────────────────────
    _set_run(paras[19], 1, f' {fd.get("kol_zhidk_tolshina", "")} ')
    _set_run(paras[20], 0, f'  Эхогенность {fd.get("kol_zhidk_exogen", "")}')
    _set_run(paras[21], 2, f' {fd.get("kol_sus_shel", "")} ')

    # ── Гиалиновый хрящ ───────────────────────────────────────────────────────
    _set_run(paras[23], 0, f'  Контур {fd.get("kol_hyalin_kontur", "")}')
    _set_run(paras[24], 1, f' {fd.get("kol_hyalin_tolshina", "")} ')
    _set_run(paras[25], 1, f'уктура {fd.get("kol_hyalin_exostr", "")}')
    _set_run(paras[26], 0, f'  Эхогенность {fd.get("kol_hyalin_exogen", "")}')

    # ── Медиальный мениск — Передний+Задний рог (side-by-side) ───────────────
    _set_run(paras[30], 0, ' Форма')
    _set_run(paras[30], 1, f' {fd.get("kol_med_per_forma", "")}')
    _set_run(paras[30], 6, ' Форма')
    _set_run(paras[30], 7, f' {fd.get("kol_med_zad_forma", "")}')
    _set_run(paras[31], 1, f' Контур {fd.get("kol_med_per_kontur", "")}')
    _set_run(paras[31], 5, f'                    Контур {fd.get("kol_med_zad_kontur", "")}')
    _set_run(paras[32], 1, f'труктура {fd.get("kol_med_per_exostr", "")} ')
    _set_run(paras[32], 9, f'Эхоструктура {fd.get("kol_med_zad_exostr", "")} ')
    _set_run(paras[33], 0, f'  Эхогенность {fd.get("kol_med_per_exogen", "")}')
    _set_run(paras[33], 8, f' Эхогенность {fd.get("kol_med_zad_exogen", "")}')

    # ── Латеральный мениск — Передний+Задний рог (side-by-side) ──────────────
    _set_run(paras[36], 0, ' Форма')
    _set_run(paras[36], 1, f' {fd.get("kol_lat_per_forma", "")}')
    _set_run(paras[36], 5, ' Форма')
    _set_run(paras[36], 6, f' {fd.get("kol_lat_zad_forma", "")}')
    _set_run(paras[37], 0, f'  Контур {fd.get("kol_lat_per_kontur", "")}')
    _set_run(paras[37], 8, f'Контур {fd.get("kol_lat_zad_kontur", "")}')
    _set_run(paras[38], 0, f'  Эхоструктура {fd.get("kol_lat_per_exostr", "")} ')
    _set_run(paras[38], 8, f'Эхоструктура {fd.get("kol_lat_zad_exostr", "")} ')
    _set_run(paras[39], 0, f'  Эхогенность {fd.get("kol_lat_per_exogen", "")}')
    _set_run(paras[39], 10, f'  Эхогенность {fd.get("kol_lat_zad_exogen", "")} ')

    # ── Костные структуры ─────────────────────────────────────────────────────
    _set_run(paras[41], 0, f'Бедренная кость {fd.get("kol_bedro", "")}')
    _set_run(paras[42], 0, f'Большеберцовая кость {fd.get("kol_bolshebert", "")}')
    _set_run(paras[43], 0, f'Надколенник {fd.get("kol_nadkolenn", "")}')

    # [47] Заключение: r[0]='З', r[1]='аключение.' — must clear r[1]
    _set_run(paras[47], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_run(paras[47], 1, '')
    _set_recom(paras[48], fd.get('rekomendacii', ''))


# ─── Protocol 12: I триместр беременности ────────────────────────────────────

def _fill_protocol_12(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # [8] Дата менструации: r[2]='  ,', r[3]=' соответствует  ', r[4]='нед', r[5]='.  дням'
    _set_run(paras[8], 2, f' {fd.get("tri1_data_mens", "")},')
    _set_run(paras[8], 3, f' соответствует {fd.get("tri1_srok_mens", "")} ')
    _set_run(paras[8], 4, '')
    _set_run(paras[8], 5, '')

    # [9] Дата родов: r[2]='  ', r[4]=' по размера', r[5]='м '
    _set_run(paras[9], 2, f' {fd.get("tri1_data_rodov_mens", "")} ')
    _set_run(paras[9], 4, f' по размерам {fd.get("tri1_data_rodov_razm", "")}')
    _set_run(paras[9], 5, '')

    # ── Матка ─────────────────────────────────────────────────────────────────
    _set_run(paras[11], 1, f'ур {fd.get("tri1_mat_kontur", "")}')
    _set_run(paras[12], 3, f' {fd.get("tri1_mat_exostr", "")}')

    # [14] Размеры матки: r[0]=dlina, r[5]=tolshina blank, r[11]=shirina blank
    _set_run(paras[14], 0, f'Длина {fd.get("tri1_mat_dlina", "")} ')
    _set_run(paras[14], 5, f' {fd.get("tri1_mat_tolshina", "")} ')
    _set_run(paras[14], 11, f' {fd.get("tri1_mat_shirina", "")} ')

    _set_run(paras[16], 1, f'и: {fd.get("tri1_mat_osoben", "")}.')

    # ── Плодное яйцо ──────────────────────────────────────────────────────────
    _set_run(paras[18], 0, f'Количество {fd.get("tri1_plod_kol", "1")}')
    _set_run(paras[19], 1, f'опография {fd.get("tri1_plod_topogr", "")}')
    _set_run(paras[20], 1, f'ма {fd.get("tri1_plod_forma", "")}')

    # [21] Диаметр/срок: r[0]=diam, r[2]=srok, r[3-5]=clear
    _set_run(paras[21], 0, f'Средний диаметр {fd.get("tri1_plod_diam", "")} ')
    _set_run(paras[21], 2, f', соответствует {fd.get("tri1_plod_srok", "")} ')
    _set_run(paras[21], 3, '')
    _set_run(paras[21], 4, '')
    _set_run(paras[21], 5, '')

    # ── Эмбрион ───────────────────────────────────────────────────────────────
    _set_run(paras[24], 0, f'Количество {fd.get("tri1_emb_kol", "1")}.')

    # [25] КТР/срок: r[0]=ktr, r[2]=srok, r[4]='.'
    _set_run(paras[25], 0, f'КТР эмбриона {fd.get("tri1_emb_ktr", "")} ')
    _set_run(paras[25], 2, f', соответствует {fd.get("tri1_emb_ktr_srok", "")} ')
    _set_run(paras[25], 4, '.')

    # [26] Сердцебиение/ЧСС: r[0] ends with blank for chss, r[1]='ударов в минуту.'
    _set_run(paras[26], 0,
             f'Сердцебиение {fd.get("tri1_emb_serd", "")},   {fd.get("tri1_emb_chss", "")} ')

    # [28] Воротниковое пространство: r[2]='  '(blank), r[3]=' ', r[4]='мм'
    _set_run(paras[28], 2, f' {fd.get("tri1_emb_vp", "")} ')

    # [30] Желточный мешочек: r[2]='визуализируется'
    _set_run(paras[30], 2, fd.get('tri1_emb_zhel', ''))

    # [31] Диаметр: r[1]=' '
    _set_run(paras[31], 1, f' {fd.get("tri1_emb_zhel_diam", "")} ')

    # ── Хорион ────────────────────────────────────────────────────────────────
    _set_run(paras[32], 1, f'Локализация {fd.get("tri1_hor_lok", "")} ')
    _set_run(paras[33], 1, f' {fd.get("tri1_hor_t", "")} ')

    # [34] Особенности — clear default
    _set_run(paras[34], 0, f'Особенности: {fd.get("tri1_osobennosti", "")}')
    # ── Яичники ───────────────────────────────────────────────────────────────
    _set_run(paras[36], 2, fd.get('tri1_yachniki', ''))

    _set_run(paras[37], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_recom(paras[38], fd.get('rekomendacii', ''))


# ─── Protocol 13: Плод ────────────────────────────────────────────────────────

def _fill_protocol_13(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # [9] Дата менструации: r[2]=' ,', r[4]='т сроку беременности  ', r[5]='нед', r[6]='.  д', r[7]='.'
    _set_run(paras[9], 2, f' {fd.get("plod_data_mens", "")},')
    _set_run(paras[9], 4, f'т сроку беременности {fd.get("plod_srok", "")} ')
    _set_run(paras[9], 5, '')
    _set_run(paras[9], 6, '')
    _set_run(paras[9], 7, '')

    # [10] Дата родов: r[2]='  ', r[5]='  '
    _set_run(paras[10], 2, f' {fd.get("plod_data_rodov_mens", "")} ')
    _set_run(paras[10], 5, f' {fd.get("plod_data_rodov_razm", "")} ')

    # ── Матка ─────────────────────────────────────────────────────────────────
    _set_run(paras[12], 1, f'тур {fd.get("plod_mat_kontur", "")} ')
    mat_exostr = fd.get('plod_mat_exostr', '')
    mat_osob = fd.get('plod_mat_osob', '').strip() if fd.get('plod_mat_osob') else ''
    osob_suffix = f' Особенности {mat_osob}.' if mat_osob else ''
    _set_run(paras[13], 2, f' {mat_exostr}{osob_suffix} ')

    # ── Плод ──────────────────────────────────────────────────────────────────
    _set_run(paras[16], 1, f'ичество {fd.get("plod_kol", "")}.')
    _set_run(paras[17], 1, f' {fd.get("plod_polozhenie", "")}')
    # [18] Движения: r[1]=туловища, r[5]=конечностями, r[8]=глотательные, r[13]=сосательные
    _set_run(paras[18], 1, f' {fd.get("plod_dvizh_tul", "")} ')
    _set_run(paras[18], 5, f' {fd.get("plod_dvizh_konech", "")} ')
    _set_run(paras[18], 8, f' {fd.get("plod_dvizh_glotat", "")} ')
    _set_run(paras[18], 13, f' {fd.get("plod_dvizh_sosat", "")} ')

    # [19] Сердцебиение/ЧСС: r[0]=label+serd, r[1]=chss+'ударов...'
    _set_run(paras[19], 0, f'Сердцебиение {fd.get("plod_serd", "")}, ')
    _set_run(paras[19], 1, f'{fd.get("plod_chss", "")} ударов в 1 минуту ')

    # ── Фетометрия ────────────────────────────────────────────────────────────
    # [21] БПР: r[1]=bpd, r[5]=srok, r[7-8]=clear
    _set_run(paras[21], 1, f'ный размер {fd.get("plod_bpd", "")} ')
    _set_run(paras[21], 5, f' {fd.get("plod_bpd_srok", "")} ')
    _set_run(paras[21], 7, '')
    _set_run(paras[21], 8, '')

    # [22] ЛЗР: r[1]=lor, r[5]=srok, r[6-7]=clear
    _set_run(paras[22], 1, f'чный размер {fd.get("plod_lor", "")} ')
    _set_run(paras[22], 5, f' {fd.get("plod_lor_srok", "")} ')
    _set_run(paras[22], 6, '')
    _set_run(paras[22], 7, '')

    # [23] Цефалический индекс: r[1]
    _set_run(paras[23], 1, f'ефалический индекс {fd.get("plod_cefal", "")} ')

    # [24] Бедренная кость: r[1]=bedro, r[8]=srok, r[9]=clear
    _set_run(paras[24], 1, f'ой кости {fd.get("plod_bedro", "")} ')
    _set_run(paras[24], 8, f' {fd.get("plod_bedro_srok", "")} ')
    _set_run(paras[24], 9, '')
    _set_run(paras[24], 10, '')

    # [25] Окружность головки: r[1]=okr, r[8]=srok, r[9]=clear
    _set_run(paras[25], 1, f'ловки плода {fd.get("plod_okr_golova", "")} ')
    _set_run(paras[25], 8, f' {fd.get("plod_okr_golova_srok", "")} ')
    _set_run(paras[25], 9, '')

    # [26] Окружность живота: r[0]=okr, r[7]=srok
    _set_run(paras[26], 0, f'Окружность живота {fd.get("plod_okr_zhivot", "")} ')
    _set_run(paras[26], 7, f' {fd.get("plod_okr_zhivot_srok", "")} ')

    osobennosti = fd.get('plod_osobennosti', '')
    pupovina = fd.get('plod_pupovina', '')
    osb_text = f'Особенности плода {osobennosti}.\n' if osobennosti else ''
    _set_run(paras[27], 0, f'{osb_text}Пуповина')
    _set_run(paras[27], 1, f' {pupovina}.')

    # ── Плацента ──────────────────────────────────────────────────────────────
    _set_run(paras[29], 1, f' {fd.get("plac_lok", "")} ')
    _set_run(paras[30], 1, f' {fd.get("plac_tolshina", "")} ')
    _set_run(paras[31], 1, f' {fd.get("plac_zrelost", "")} степени.')
    # [32] Эхоструктура плаценты: r[0]="Эхоструктура", r[1]=" {val}."
    _set_run(paras[32], 0, 'Эхоструктура')
    _set_run(paras[32], 1, f' {fd.get("plac_exostr", "")}.')

    # ── Воды ──────────────────────────────────────────────────────────────────
    # [33] Околоплодные воды: r[1]=water value+comma, r[2]=ИАЖ+value, r[3-4]=clear
    _set_run(paras[33], 1, f' {fd.get("vodi_water", "")} , ')
    _set_run(paras[33], 2, f' ИАЖ {fd.get("vodi_iazh", "")} ')
    _set_run(paras[33], 3, '')
    _set_run(paras[33], 4, '')

    # [34] Эхогенность: r[0]="Эхогенность {val}"
    _set_run(paras[34], 0, f'Эхогенность {fd.get("vodi_exogen", "")} ')

    # [35] Масса плода
    _set_run(paras[35], 0, f'Масса плода {fd.get("plod_massa", "")} ')

    # [36] Яичники: r[0]="Яичники", r[1]=" value."
    _set_run(paras[36], 0, f'Яичники')
    _set_run(paras[36], 1, f' {fd.get("plod_yachniki", "")}.')

    _set_run(paras[37], 0, f'Заключение: {fd.get("zaklyuchenie", "")}')
    _set_run(paras[37], 1, '')
    _set_recom(paras[38], fd.get('rekomendacii', ''))


# ─── Protocol 14: Сердцебиение ────────────────────────────────────────────────
#
# Template (starts at para[0] with date):
#   [1] Средний диаметр: r[0]='Средний диаметр плодного яйца ', r[1]='мм', r[2]=', соответствует...', r[3]='нед', r[4]='.  д.'
#   [2] КТР: r[0]='КТР эмбриона  ', r[1]='мм', r[2]=', соответствует...', r[3]='нед', r[4]='.   д.'
#   [3] Сердцебиение: r[0]='Сердцебиение', r[1]='   ,', r[2]=' ЧСС   ударов в минуту.'
#   [4] Желточный мешочек: r[0]='Желточный мешочек диаметром  ', r[1]='мм', r[2]='.'
#   [5] Заключение: r[0]='Заключение:', r[1]=' беременность по КТР   ', r[2]='нед', r[3]='.    д.,...'

def _fill_protocol_14(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # [1] Средний диаметр: r[0] has value blank, r[1]='мм' stays
    _set_run(paras[1], 0, f'Средний диаметр плодного яйца {fd.get("serd_plod_diam", "")} ')
    _set_run(paras[1], 3, fd.get('serd_plod_srok', ''))
    _set_run(paras[1], 4, '.')

    # [2] КТР: r[0] has value blank, r[1]='мм' stays
    _set_run(paras[2], 0, f'КТР эмбриона {fd.get("serd_ktr", "")} ')
    _set_run(paras[2], 3, fd.get('serd_ktr_srok', ''))
    _set_run(paras[2], 4, '.')

    # [3] Сердцебиение/ЧСС
    _set_run(paras[3], 1, f' {fd.get("serd_serdb", "")}, ')
    _set_run(paras[3], 2, f' ЧСС {fd.get("serd_chss", "")} ударов в минуту.')

    # [4] Желточный мешочек: r[0] has value blank, r[1]='мм' stays
    _set_run(paras[4], 0, f'Желточный мешочек диаметром {fd.get("serd_zhel_diam", "")} ')

    # [5] Заключение: put full text in r[1], clear r[2] and r[3]
    _set_run(paras[5], 1, f' {fd.get("zaklyuchenie", "")}')
    _set_run(paras[5], 2, '')
    _set_run(paras[5], 3, '')


# ─── Per-protocol fill dispatch ───────────────────────────────────────────────

# ─── Protocol 15: Фолликулометрия ────────────────────────────────────────────

def _fill_protocol_15(doc: Document, fd: Dict):
    paras = doc.paragraphs

    # [9] Толщина эндометрия (М-эхо): r[1]=value
    _set_run(paras[9], 1, f' {fd.get("fol_endo_t", "")} ')

    # [10] Эндометрий соответствует фазе: r[1]=value
    _set_run(paras[10], 1, f' {fd.get("fol_endo_faza", "")} ')

    # [16] Правый яичник — фолликулы: r[1]=value
    _set_run(paras[16], 1, f' {fd.get("fol_pr_follikuli", "")} ')

    # [19] Левый яичник — фолликулы: r[1]=value
    _set_run(paras[19], 1, f' {fd.get("fol_lev_follikuli", "")} ')

    # [21] Свободная жидкость: r[1]=value
    _set_run(paras[21], 1, f' {fd.get("fol_svobod_zhid", "")} ')

    # [23] Заключение: r[1]=value
    _set_run(paras[23], 1, f' {fd.get("zaklyuchenie", "")} ')


# ─── Protocol 16: Мошонка (shablonsiz, to'liq yaratiladi) ────────────────────

def _build_protocol_16(form_data: Dict, client: Dict, doctor: Dict, created_at: str) -> 'Document':
    from docx import Document as DocxDoc
    doc = DocxDoc()

    # Marginlar
    for section in doc.sections:
        section.top_margin = Mm(1)
        section.bottom_margin = Mm(1)
        section.left_margin = Mm(1)
        section.right_margin = Mm(1)

    fd = form_data

    def _p(text='', bold=False, center=False, right=False, space_after=0):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(16)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if text:
            r = p.add_run(text)
            r.font.size = Pt(12)
            r.bold = bold
        return p

    def _val(key):
        v = str(fd.get(key, '') or '').strip()
        return f'  {v}  ' if v else '   '

    # Sarlavha
    _p('Медицинский центр «NEVROCARDIOMED»', bold=True, center=True, space_after=10)

    # Bemor va sana
    last = client.get('last_name', '') or ''
    first = client.get('first_name', '') or ''
    patronymic = client.get('patronymic', '') or ''
    full_name = ' '.join(p for p in [last, first, patronymic] if p)
    birth_year = str(client.get('birth_date', '') or '')[:4]
    _p(f'Ф.И.О. пациента {full_name}, {birth_year} г.р.')
    _p(f'Дата проведения исследования: {(created_at or "")[:10]}')
    _p('Ультразвуковой аппарат SonoScape S50. Тип датчика – линейный, 7,5-12,0 МГц.', space_after=10)

    # Protokol sarlavhasi
    _p('ПРОТОКОЛ', bold=True, center=True)
    _p('ультразвукового исследования органов мошонки.', center=True, space_after=6)

    # Yaichki — Pravoe
    _p('ЯИЧКИ', bold=True)
    _p('Правое', bold=True)
    _p(f'Топография{_val("yai_pr_topografiya")}.')
    _p(f'Форма{_val("yai_pr_forma")}.')
    _p(f'Контур{_val("yai_pr_kontur")}(в норме ровный, четкий).')
    _p(f'Капсула (белочная оболочка){_val("yai_pr_kapsula")}.')
    _p('Размеры линейные:')
    _p(f'длина{_val("yai_pr_dlina")}мм (в норме 35-60 мм), толщина{_val("yai_pr_tolshina")}мм (в норме 15-30 мм), ширина{_val("yai_pr_shirina")}мм (в норме 15-30 мм).')
    _p(f'Объём{_val("yai_pr_obyom")}см³ (в норме 15 см³).')
    _p(f'Эхоструктура{_val("yai_pr_exostruktura")}(в норме однородная).')
    _p(f'Эхогенность{_val("yai_pr_exogennost")}(в норме средняя).')
    _p(f'Влагалище, толщина жидкости в нём{_val("yai_pr_vlagalishe")}мм.')
    _p(f'Оболочки, толщина{_val("yai_pr_obolochki")}мм (в норме 2-7 мм).')
    _p(f'Наличие расширения вен семенного канатика{_val("yai_pr_rasshireniya")}мм.')

    # Yaichki — Levoe
    _p('Левое', bold=True, space_after=2)
    _p(f'Топография{_val("yai_le_topografiya")}.')
    _p(f'Форма{_val("yai_le_forma")}.')
    _p(f'Контур{_val("yai_le_kontur")}(в норме ровный, четкий).')
    _p(f'Капсула (белочная оболочка){_val("yai_le_kapsula")}.')
    _p('Размеры линейные:')
    _p(f'длина{_val("yai_le_dlina")}мм (в норме 35-60 мм), толщина{_val("yai_le_tolshina")}мм (в норме 15-30 мм), ширина{_val("yai_le_shirina")}мм (в норме 15-30 мм).')
    _p(f'Объём{_val("yai_le_obyom")}см³ (в норме 15 см³).')
    _p(f'Эхоструктура{_val("yai_le_exostruktura")}(в норме однородная).')
    _p(f'Эхогенность{_val("yai_le_exogennost")}(в норме средняя).')
    _p(f'Влагалище, толщина жидкости в нём{_val("yai_le_vlagalishe")}мм.')
    _p(f'Оболочки, толщина{_val("yai_le_obolochki")}мм (в норме 2-7 мм).')
    _p(f'Наличие расширения вен семенного канатика{_val("yai_le_rasshireniya")}мм.')

    # Pridatki
    _p('ПРИДАТКИ', bold=True, space_after=2)
    _p('Правый', bold=True)
    _p(f'Топография{_val("pri_pr_topografiya")}.')
    _p(f'Контур{_val("pri_pr_kontur")}(в норме ровный, четкий).')
    _p('Размер, толщина:')
    _p(f'в области головки{_val("pri_pr_golovki")}мм, в области тела{_val("pri_pr_tela")}мм.')
    _p(f'Эхоструктура{_val("pri_pr_exostruktura")}(в норме однородная).')
    _p(f'Эхогенность{_val("pri_pr_exogennost")}(в норме средняя).')

    _p('Левый', bold=True, space_after=2)
    _p(f'Топография{_val("pri_le_topografiya")}.')
    _p(f'Контур{_val("pri_le_kontur")}(в норме ровный, четкий).')
    _p('Размер, толщина:')
    _p(f'в области головки{_val("pri_le_golovki")}мм, в области тела{_val("pri_le_tela")}мм.')
    _p(f'Эхоструктура{_val("pri_le_exostruktura")}(в норме однородная).')
    _p(f'Эхогенность{_val("pri_le_exogennost")}(в норме средняя).')

    # Xulosa
    zakl = (fd.get('zaklyuchenie', '') or '').strip()
    rekom = (fd.get('rekomendacii', '') or '').strip()
    _p(f'Заключение: {zakl}', space_after=0)
    if rekom:
        _p(f'Рекомендации: {rekom}')

    # Shifokor
    d_last = doctor.get('last_name', '') or ''
    d_first = doctor.get('first_name', '') or ''
    d_pat = doctor.get('patronymic_name', doctor.get('patronymic', '')) or ''
    doctor_name = ' '.join(p for p in [d_last, d_first, d_pat] if p)
    vp = _p(f'Врач_________________ {doctor_name}.', right=True, space_after=0)
    vp.paragraph_format.space_before = Pt(20)
    _p('Тел.: (+99893)7240070', right=True)

    return doc


PROTOCOL_FILLERS = {
    1:  _fill_protocol_1,
    2:  _fill_protocol_2,
    3:  _fill_protocol_3,
    4:  _fill_protocol_4,
    5:  _fill_protocol_5,
    6:  _fill_protocol_6,
    7:  _fill_protocol_7,
    8:  _fill_protocol_8,
    9:  _fill_protocol_9,
    10: _fill_protocol_10,
    11: _fill_protocol_11,
    12: _fill_protocol_12,
    13: _fill_protocol_13,
    14: _fill_protocol_14,
    15: _fill_protocol_15,
}


def generate_protocol_docx(
    protocol_id: int,
    protocol_title: str,
    form_data: Dict[str, Any],
    client: Dict[str, Any],
    doctor: Dict[str, Any],
    created_at: str = "",
) -> bytes:
    # Protocol 16 (Мошонка) — shablonsiz dinamik yaratiladi
    if protocol_id == 16:
        doc = _build_protocol_16(form_data, client, doctor, created_at)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    template_file = TEMPLATE_FILES.get(protocol_id)
    if not template_file:
        raise ValueError(f"No template for protocol_id={protocol_id}")

    template_path = os.path.join(TEMPLATES_DIR, template_file)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    for section in doc.sections:
        section.top_margin = Mm(1)
        section.bottom_margin = Mm(1)
        section.left_margin = Mm(1)
        section.right_margin = Mm(1)

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Matnli paragraf — ixcham
            para.paragraph_format.line_spacing = Pt(16)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            # Klinika nomidan keyin ko'rinadigan bo'sh joy (2 satr)
            if 'NEVROCARDIOMED' in text:
                para.paragraph_format.space_after = Pt(32)
        else:
            # Bo'sh paragraf — ko'rinadigan ajratuvchi
            para.paragraph_format.line_spacing = Pt(18)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

    # Fill common header (patient info, date, doctor)
    _fill_header(doc, client, doctor, created_at)

    # Har bir qiymatning ikki tomoniga 2 ta bo'sh joy qo'shish
    padded_data = {}
    for k, v in form_data.items():
        s = str(v).strip() if v else ''
        padded_data[k] = f'  {s}  ' if s else ''

    # Fill protocol-specific fields
    filler = PROTOCOL_FILLERS.get(protocol_id)
    if filler:
        filler(doc, padded_data)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
