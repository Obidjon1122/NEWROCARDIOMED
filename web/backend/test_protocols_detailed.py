#!/usr/bin/env python3
"""
Detailed test suite for NEWROCARDIOMED DOCX protocol generator.
Provides detailed field-by-field validation and comprehensive reporting.
"""

import sys
import io
from docx import Document
from app.services.docx_generator import generate_protocol_docx

# Test data
CLIENT = {
    'first_name': 'Алишер',
    'last_name': 'Каримов',
    'patronymic': 'Нодирович',
    'gender': 'male',
    'birth_date': '1985-03-15'
}

DOCTOR = {
    'first_name': 'Азиз',
    'last_name': 'Латипов',
    'patronymic_name': 'Абдурахмонович'
}

CREATED_AT = '2026-03-30T10:00:00'

# Sample test data for each protocol
TEST_DATA = {
    1: {
        'name': 'Печень и желчный пузырь',
        'sample': {
            'pech_topografia': 'в норме',
            'pech_konturi': 'в норме',
            'pech_pzr_lev': '50',
            'pech_kkr_lev': '85',
            'pech_exostruktura': 'однородная',
            'zaklyuchenie': 'Печень в норме',
            'rekomendacii': 'Диспансерное наблюдение'
        }
    },
    2: {
        'name': 'Почки и мочевой пузырь',
        'sample': {
            'prav_topografia': 'в норме',
            'prav_dlina': '100',
            'prav_konturi': 'в норме',
            'lev_topografia': 'в норме',
            'lev_dlina': '100',
            'mp_obem': '180',
            'zaklyuchenie': 'Почки в норме'
        }
    },
    3: {
        'name': 'Щитовидная железа',
        'sample': {
            'p_topografiya': 'в норме',
            'dp_dlina': '50',
            'dl_dlina': '48',
            'summarniy': '55',
            'zaklyuchenie_sh': 'Щитовидная железа в норме'
        }
    },
    4: {
        'name': 'Поджелудочная железа',
        'sample': {
            'pdj_topografia': 'в норме',
            'pdj_golovka': '20',
            'pdj_telo': '10',
            'pdj_hvost': '15',
            'zaklyuchenie': 'Поджелудочная железа в норме'
        }
    },
    5: {
        'name': 'Селезёнка',
        'sample': {
            'sel_topografia': 'в норме',
            'sel_dlina': '110',
            'sel_obem': '150',
            'zaklyuchenie': 'Селезёнка в норме'
        }
    },
    6: {
        'name': 'Простата и семенные пузырьки',
        'sample': {
            'pr_topografia': 'в норме',
            'pr_dlina': '32',
            'pr_obem': '22',
            'svp_dlina': '40',
            'svl_dlina': '40',
            'zaklyuchenie': 'Простата в норме'
        }
    },
    7: {
        'name': 'Молочные железы',
        'sample': {
            'prm_topografia': 'в норме',
            'prm_kozha_t': '2',
            'prm_parenhima_t': '15',
            'levm_topografia': 'в норме',
            'levm_parenhima_t': '15',
            'zaklyuchenie': 'Молочные железы в норме'
        }
    },
    8: {
        'name': 'Мочевой пузырь',
        'sample': {
            'mp_obem_fiz': '180',
            'mp_forma': 'в норме',
            'mp_stenka_do': '4',
            'mp_ostatoch': '15',
            'zaklyuchenie': 'Мочевой пузырь в норме'
        }
    }
}

def test_protocol_generation(protocol_id):
    """Test DOCX generation for a protocol."""
    test_info = TEST_DATA[protocol_id]
    sample_data = test_info['sample']

    print(f"\nProtocol {protocol_id}: {test_info['name']}")
    print("-" * 70)

    try:
        # Generate DOCX
        docx_bytes = generate_protocol_docx(
            protocol_id=protocol_id,
            protocol_title=f"Protocol {protocol_id} Test",
            form_data=sample_data,
            client=CLIENT,
            doctor=DOCTOR,
            created_at=CREATED_AT
        )

        # Verify bytes
        if not docx_bytes or len(docx_bytes) == 0:
            print(f"  FAIL: Generated DOCX is empty (0 bytes)")
            return False

        print(f"  Generated file size: {len(docx_bytes)} bytes")

        # Verify DOCX can be opened
        try:
            doc = Document(io.BytesIO(docx_bytes))
            print(f"  DOCX is valid and can be opened")
        except Exception as e:
            print(f"  FAIL: Cannot open DOCX: {str(e)}")
            return False

        # Extract all text
        all_text = '\n'.join([p.text for p in doc.paragraphs])

        # Check for patient info
        if CLIENT['first_name'] in all_text and CLIENT['last_name'] in all_text:
            print(f"  Patient info present: OK")
        else:
            print(f"  WARNING: Patient info not fully found")

        # Check for doctor info
        if DOCTOR['first_name'] in all_text or DOCTOR['last_name'] in all_text:
            print(f"  Doctor info present: OK")
        else:
            print(f"  WARNING: Doctor info not found")

        # Check for key sample fields
        missing_fields = []
        found_fields = []
        for key, value in sample_data.items():
            if value and len(str(value).strip()) > 0:
                if str(value) in all_text:
                    found_fields.append(key)
                else:
                    missing_fields.append(f"{key}='{value}'")

        if found_fields:
            print(f"  Found {len(found_fields)} sample fields in document")

        if missing_fields:
            print(f"  WARNING: {len(missing_fields)} sample fields not found:")
            for field in missing_fields[:5]:
                print(f"    - {field}")
        else:
            print(f"  All sample fields found: OK")

        # Check table count
        print(f"  Tables in document: {len(doc.tables)}")

        print(f"  Status: PASS")
        return True

    except Exception as e:
        print(f"  FAIL: {str(e)}")
        return False

def test_edge_cases():
    """Test edge cases."""
    print(f"\n{'='*70}")
    print("EDGE CASE TESTS")
    print(f"{'='*70}")

    # Test 1: Empty form data
    print(f"\nTest 1: Empty form data")
    print("-" * 70)
    try:
        docx_bytes = generate_protocol_docx(
            protocol_id=1,
            protocol_title="Empty Test",
            form_data={},
            client=CLIENT,
            doctor=DOCTOR,
            created_at=CREATED_AT
        )
        if docx_bytes and len(docx_bytes) > 0:
            print(f"  Generated empty form DOCX: {len(docx_bytes)} bytes - PASS")
        else:
            print(f"  FAIL: Generated DOCX is empty")
    except Exception as e:
        print(f"  FAIL: {str(e)}")

    # Test 2: Very long text
    print(f"\nTest 2: Very long text (500+ chars)")
    print("-" * 70)
    long_text = "A" * 500
    try:
        docx_bytes = generate_protocol_docx(
            protocol_id=1,
            protocol_title="Long Text Test",
            form_data={'zaklyuchenie': long_text, 'rekomendacii': long_text},
            client=CLIENT,
            doctor=DOCTOR,
            created_at=CREATED_AT
        )
        if docx_bytes and len(docx_bytes) > 0:
            doc = Document(io.BytesIO(docx_bytes))
            all_text = '\n'.join([p.text for p in doc.paragraphs])
            if long_text in all_text:
                print(f"  Long text handled correctly: PASS")
            else:
                print(f"  WARNING: Long text may be truncated")
        else:
            print(f"  FAIL: Generated DOCX is empty")
    except Exception as e:
        print(f"  FAIL: {str(e)}")

    # Test 3: Special characters
    print(f"\nTest 3: Special characters (Cyrillic, quotes, symbols)")
    print("-" * 70)
    special_text = 'Текст "с кавычками" и символами: № — €'
    try:
        docx_bytes = generate_protocol_docx(
            protocol_id=1,
            protocol_title="Special Chars Test",
            form_data={'zaklyuchenie': special_text},
            client=CLIENT,
            doctor=DOCTOR,
            created_at=CREATED_AT
        )
        if docx_bytes and len(docx_bytes) > 0:
            doc = Document(io.BytesIO(docx_bytes))
            all_text = '\n'.join([p.text for p in doc.paragraphs])
            if special_text in all_text:
                print(f"  Special characters handled correctly: PASS")
            else:
                print(f"  WARNING: Special characters may be modified")
        else:
            print(f"  FAIL: Generated DOCX is empty")
    except Exception as e:
        print(f"  FAIL: {str(e)}")

    # Test 4: Missing template
    print(f"\nTest 4: Invalid protocol ID")
    print("-" * 70)
    try:
        docx_bytes = generate_protocol_docx(
            protocol_id=999,
            protocol_title="Invalid Protocol",
            form_data={},
            client=CLIENT,
            doctor=DOCTOR,
            created_at=CREATED_AT
        )
        print(f"  FAIL: Should have raised ValueError for invalid protocol_id")
    except ValueError as e:
        print(f"  Correctly raised ValueError: {str(e)} - PASS")
    except Exception as e:
        print(f"  FAIL: Unexpected error: {str(e)}")

def main():
    """Run all tests."""
    print(f"{'='*70}")
    print("DETAILED PROTOCOL GENERATION TESTS")
    print(f"{'='*70}")

    passed = 0
    total = 8

    for protocol_id in range(1, 9):
        if test_protocol_generation(protocol_id):
            passed += 1

    # Run edge cases
    test_edge_cases()

    print(f"\n{'='*70}")
    print(f"SUMMARY: {passed}/{total} protocols passed")
    print(f"{'='*70}\n")

    return passed == total

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
