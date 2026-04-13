#!/usr/bin/env python3
"""
Comprehensive test suite for NEWROCARDIOMED DOCX protocol generator.
Tests protocols 1-8 with 5 scenarios each:
1. Normal/healthy values (all defaults)
2. All pathological/abnormal options selected (combobox fields set to non-default values)
3. All fields empty (empty strings)
4. Very long text in textarea fields (200+ chars)
5. Mixed scenario (some fields filled, some empty)
"""

import sys
import io
from docx import Document
from app.services.docx_generator import generate_protocol_docx

# Test data
CLIENT = {
    'first_name': 'Алишер',
    'last_name': 'Каримов',
    'patronymic': 'Нодирович',
    'gender': 'male',
    'birth_date': '1985-03-15'
}

DOCTOR = {
    'first_name': 'Азиз',
    'last_name': 'Латипов',
    'patronymic_name': 'Абдурахмонович'
}

CREATED_AT = '2026-03-30T10:00:00'

# ────────────────────────────────────────────────────────────────────────────
# PROTOCOL FIELD DEFINITIONS FOR EACH PROTOCOL 1-8
# ────────────────────────────────────────────────────────────────────────────

PROTOCOL_1_FIELDS = {
    'pech_topografia', 'pech_konturi', 'pech_pzr_lev', 'pech_kkr_lev',
    'pech_pzr_hvost', 'pech_pzr_prav', 'pech_kkr_prav', 'pech_kvr',
    'pech_lat_razmer', 'pech_ugol_prav', 'pech_ugol_lev', 'pech_obem',
    'pech_massa', 'pech_exostruktura', 'pech_exogennost', 'pech_zvukoprov',
    'soud_pech_veny', 'soud_sum_diam', 'soud_port_vena', 'soud_nizhnyaya_pv',
    'zhp_vnutrip', 'zhp_obshchiy', 'puz_topografia', 'puz_forma',
    'puz_kontur', 'puz_dlina', 'puz_tolshina', 'puz_shirina', 'puz_obem',
    'puz_ploshad', 'puz_protok', 'puz_stenka', 'puz_soderzhimoe',
    'zaklyuchenie', 'rekomendacii'
}

PROTOCOL_2_FIELDS = {
    'prav_topografia', 'prav_podvizhnost', 'prav_zhir_kapsul', 'prav_fib_kapsul',
    'prav_konturi', 'prav_dlina', 'prav_tolshina', 'prav_shirina', 'prav_obem',
    'prav_parenhima_t', 'prav_exogen_par', 'prav_exostr_par', 'prav_chls_tolshina',
    'prav_chls_exostr', 'prav_kri', 'prav_konkrement', 'prav_mochetchnik',
    'lev_topografia', 'lev_podvizhnost', 'lev_zhir_kapsul', 'lev_fib_kapsul',
    'lev_konturi', 'lev_dlina', 'lev_tolshina', 'lev_shirina', 'lev_obem',
    'lev_parenhima_t', 'lev_exogen_par', 'lev_exostr_par', 'lev_chls_tolshina',
    'lev_chls_exostr', 'lev_kri', 'lev_konkrement', 'lev_mochetchnik',
    'mp_obem', 'mp_forma', 'mp_kontur_naruzh', 'mp_kontur_vnutr',
    'mp_stenka_do', 'mp_stenka_posle', 'mp_soderzhimoe', 'mp_sheika',
    'mp_ustya', 'mp_ostatoch', 'zaklyuchenie'
}

PROTOCOL_3_FIELDS = {
    'p_topografiya', 'p_kontur', 'p_elastichnost', 'p_podvijnost', 'p_razmer',
    'p_ehostruktura', 'p_ehogennost',
    'dp_topografiya', 'dp_kontur', 'dp_elastichnost', 'dp_podvijnost', 'dp_dlina',
    'dp_tolshina', 'dp_shirina', 'dp_obem', 'dp_forma', 'dp_ehostruktura',
    'dp_ehogennost',
    'dl_topografiya', 'dl_kontur', 'dl_elastichnost', 'dl_podvijnost', 'dl_dlina',
    'dl_tolshina', 'dl_shirina', 'dl_obem', 'dl_forma', 'dl_ehostruktura',
    'dl_ehogennost',
    'summarniy', 'kg', 'sm',
    'zaklyuchenie_sh', 'rekomendatsi_sh', 'vrach', 'telefon'
}

PROTOCOL_4_FIELDS = {
    'pdj_topografia', 'pdj_forma', 'pdj_kontur', 'pdj_golovka', 'pdj_telo',
    'pdj_hvost', 'pdj_dlina', 'pdj_exostruktura', 'pdj_exogennost',
    'pdj_zvukoprov', 'pdj_virsungov_g', 'pdj_virsungov_t',
    'zaklyuchenie', 'rekomendacii'
}

PROTOCOL_5_FIELDS = {
    'sel_topografia', 'sel_forma', 'sel_konturi', 'sel_dlina', 'sel_tolshina',
    'sel_shirina', 'sel_obem', 'sel_indeks', 'sel_exostruktura', 'sel_exogennost',
    'sel_zvukoprov', 'sel_lienalis', 'zaklyuchenie', 'rekomendacii'
}

PROTOCOL_6_FIELDS = {
    'pr_topografia', 'pr_forma', 'pr_kontur', 'pr_kapsul_t', 'pr_dlina',
    'pr_tolshina', 'pr_shirina', 'pr_obem', 'pr_indeks', 'pr_exostruktura',
    'pr_exogennost', 'pr_urethra', 'pr_nach_obem', 'pr_ostatoch',
    'svp_topografia', 'svp_kontur', 'svp_dlina', 'svp_tolshina', 'svp_shirina',
    'svp_exostr_do', 'svp_exostr_posle', 'svp_exogennost',
    'svl_topografia', 'svl_kontur', 'svl_dlina', 'svl_tolshina', 'svl_shirina',
    'svl_exostr_do', 'svl_exostr_posle', 'svl_exogennost',
    'zaklyuchenie', 'rekomendacii'
}

PROTOCOL_7_FIELDS = {
    'prm_topografia', 'prm_kozha_kontur', 'prm_kozha_t', 'prm_kozha_exogen',
    'prm_kozha_exostr', 'prm_sosok', 'prm_podkozh_t', 'prm_intramam',
    'prm_retromam_t', 'prm_parenhima_kontur', 'prm_parenhima_t', 'prm_parenhima_exogen',
    'prm_parenhima_exostr', 'prm_galaktofori', 'prm_limfuzli',
    'levm_topografia', 'levm_kozha_kontur', 'levm_kozha_t', 'levm_kozha_exogen',
    'levm_kozha_exostr', 'levm_sosok', 'levm_podkozh_t', 'levm_intramam',
    'levm_retromam_t', 'levm_parenhima_kontur', 'levm_parenhima_t', 'levm_parenhima_exogen',
    'levm_parenhima_exostr', 'levm_galaktofori', 'levm_limfuzli',
    'zaklyuchenie', 'rekomendacii'
}

PROTOCOL_8_FIELDS = {
    'mp_obem_fiz', 'mp_forma', 'mp_kontur_naruzh', 'mp_kontur_vnutr',
    'mp_stenka_do', 'mp_stenka_posle', 'mp_soderzhimoe', 'mp_sheika',
    'mp_ustya', 'mp_ostatoch', 'zaklyuchenie', 'rekomendacii'
}

PROTOCOL_FIELDS_MAP = {
    1: PROTOCOL_1_FIELDS,
    2: PROTOCOL_2_FIELDS,
    3: PROTOCOL_3_FIELDS,
    4: PROTOCOL_4_FIELDS,
    5: PROTOCOL_5_FIELDS,
    6: PROTOCOL_6_FIELDS,
    7: PROTOCOL_7_FIELDS,
    8: PROTOCOL_8_FIELDS,
}

# ────────────────────────────────────────────────────────────────────────────
# SCENARIO 1: Normal/healthy values (all defaults)
# ────────────────────────────────────────────────────────────────────────────

SCENARIO_1_DATA = {
    1: {
        'pech_topografia': 'в норме',
        'pech_konturi': 'в норме',
        'pech_exostruktura': 'однородная',
        'pech_exogennost': 'средняя',
        'pech_zvukoprov': 'сохранена',
        'puz_topografia': 'в норме',
        'puz_forma': 'грушевидная',
        'puz_kontur': 'в норме',
        'zhp_vnutrip': 'не визуализируются',
        'puz_protok': 'не визуализируется',
        'puz_soderzhimoe': 'анэхогенное',
        'pech_pzr_lev': '50',
        'pech_kkr_lev': '85',
        'pech_pzr_hvost': '27',
        'pech_pzr_prav': '110',
        'pech_kkr_prav': '130',
        'pech_kvr': '20',
        'pech_lat_razmer': '21',
        'pech_obem': '1500',
        'pech_massa': '1400',
        'soud_pech_veny': '8',
        'soud_sum_diam': '29',
        'soud_port_vena': '12',
        'soud_nizhnyaya_pv': '22',
        'zhp_obshchiy': '5',
        'puz_dlina': '70',
        'puz_tolshina': '25',
        'puz_shirina': '35',
        'puz_obem': '25',
        'puz_ploshad': '16',
        'puz_stenka': '2',
        'pech_ugol_prav': 'в норме',
        'pech_ugol_lev': 'в норме',
        'zaklyuchenie': 'Печень в норме',
        'rekomendacii': 'Диспансерное наблюдение'
    },
    2: {
        'prav_topografia': 'в норме',
        'prav_podvizhnost': 'сохранена',
        'prav_zhir_kapsul': 'без особенностей',
        'prav_fib_kapsul': 'без особенностей',
        'prav_konturi': 'в норме',
        'prav_exogen_par': 'сохранена',
        'prav_exostr_par': 'без особенностей',
        'prav_chls_exostr': 'без особенностей',
        'prav_konkrement': 'не визуализируются',
        'prav_mochetchnik': 'без особенностей',
        'prav_dlina': '100',
        'prav_tolshina': '40',
        'prav_shirina': '55',
        'prav_obem': '110',
        'prav_parenhima_t': '15',
        'prav_chls_tolshina': '3',
        'prav_kri': '0.37',
        'lev_topografia': 'в норме',
        'lev_podvizhnost': 'сохранена',
        'lev_zhir_kapsul': 'без особенностей',
        'lev_fib_kapsul': 'без особенностей',
        'lev_konturi': 'в норме',
        'lev_exogen_par': 'сохранена',
        'lev_exostr_par': 'без особенностей',
        'lev_chls_exostr': 'без особенностей',
        'lev_konkrement': 'не визуализируются',
        'lev_mochetchnik': 'без особенностей',
        'lev_dlina': '100',
        'lev_tolshina': '40',
        'lev_shirina': '55',
        'lev_obem': '110',
        'lev_parenhima_t': '15',
        'lev_chls_tolshina': '3',
        'lev_kri': '0.37',
        'mp_obem': '180',
        'mp_forma': 'в норме',
        'mp_kontur_naruzh': 'в норме',
        'mp_kontur_vnutr': 'в норме',
        'mp_stenka_do': '4',
        'mp_stenka_posle': '6',
        'mp_soderzhimoe': 'анэхогенное',
        'mp_sheika': 'в норме',
        'mp_ustya': 'в норме',
        'mp_ostatoch': '10',
        'zaklyuchenie': 'Почки в норме'
    },
    3: {
        'p_topografiya': 'в норме',
        'p_kontur': 'в норме',
        'p_elastichnost': 'в норме',
        'p_podvijnost': 'в норме',
        'p_razmer': '5',
        'p_ehostruktura': 'однородная',
        'p_ehogennost': 'средняя',
        'dp_topografiya': 'в норме',
        'dp_kontur': 'в норме',
        'dp_elastichnost': 'сохранена',
        'dp_podvijnost': 'в норме',
        'dp_dlina': '50',
        'dp_tolshina': '20',
        'dp_shirina': '30',
        'dp_obem': '30',
        'dp_forma': 'в норме',
        'dp_ehostruktura': 'однородная',
        'dp_ehogennost': 'средняя',
        'dl_topografiya': 'в норме',
        'dl_kontur': 'в норме',
        'dl_elastichnost': 'в норме',
        'dl_podvijnost': 'сохранена',
        'dl_dlina': '48',
        'dl_tolshina': '18',
        'dl_shirina': '28',
        'dl_obem': '25',
        'dl_forma': 'в норме',
        'dl_ehostruktura': 'однородная',
        'dl_ehogennost': 'средняя',
        'summarniy': '55',
        'kg': '70',
        'sm': '30',
        'zaklyuchenie_sh': 'Щитовидная железа в норме',
        'rekomendatsi_sh': 'Контрольное УЗИ в год',
        'vrach': 'А.Латипов',
        'telefon': '+992931234567'
    },
    4: {
        'pdj_topografia': 'в норме',
        'pdj_forma': 'в норме',
        'pdj_kontur': 'в норме',
        'pdj_exostruktura': 'однородная',
        'pdj_exogennost': 'в норме',
        'pdj_zvukoprov': 'в норме',
        'pdj_golovka': '20',
        'pdj_telo': '10',
        'pdj_hvost': '15',
        'pdj_dlina': '100',
        'pdj_virsungov_g': '3',
        'pdj_virsungov_t': '2',
        'zaklyuchenie': 'Поджелудочная железа в норме',
        'rekomendacii': 'Здоровый образ жизни'
    },
    5: {
        'sel_topografia': 'в норме',
        'sel_forma': 'не изменена',
        'sel_konturi': 'в норме',
        'sel_exostruktura': 'в норме',
        'sel_exogennost': 'в норме',
        'sel_zvukoprov': 'в норме',
        'sel_dlina': '110',
        'sel_tolshina': '40',
        'sel_shirina': '65',
        'sel_obem': '150',
        'sel_indeks': '20',
        'sel_lienalis': '5',
        'zaklyuchenie': 'Селезёнка в норме',
        'rekomendacii': 'Диспансерное наблюдение'
    },
    6: {
        'pr_topografia': 'в норме',
        'pr_forma': 'в норме',
        'pr_kontur': 'ровный, чёткий',
        'pr_exostruktura': 'без особенностей',
        'pr_exogennost': 'без особенностей',
        'pr_urethra': 'в норме',
        'pr_kapsul_t': '1',
        'pr_dlina': '32',
        'pr_tolshina': '20',
        'pr_shirina': '35',
        'pr_obem': '22',
        'pr_indeks': '0.8',
        'pr_nach_obem': '150',
        'pr_ostatoch': '15',
        'svp_topografia': 'в норме',
        'svp_kontur': 'в норме',
        'svp_exostr_do': 'в норме',
        'svp_exostr_posle': 'в норме',
        'svp_exogennost': 'в норме',
        'svp_dlina': '40',
        'svp_tolshina': '12',
        'svp_shirina': '18',
        'svl_topografia': 'в норме',
        'svl_kontur': 'в норме',
        'svl_exostr_do': 'в норме',
        'svl_exostr_posle': 'в норме',
        'svl_exogennost': 'в норме',
        'svl_dlina': '40',
        'svl_tolshina': '12',
        'svl_shirina': '18',
        'zaklyuchenie': 'Простата в норме',
        'rekomendacii': 'Диспансерное наблюдение'
    },
    7: {
        'prm_topografia': 'в норме',
        'prm_kozha_kontur': 'в норме',
        'prm_kozha_exogen': 'в норме',
        'prm_kozha_exostr': 'в норме',
        'prm_sosok': 'без особенностей',
        'prm_intramam': 'без особенностей',
        'prm_parenhima_kontur': 'в норме',
        'prm_parenhima_exogen': 'в норме',
        'prm_parenhima_exostr': 'без особенностей',
        'prm_galaktofori': 'не визуализируются',
        'prm_limfuzli': 'без особенностей',
        'prm_kozha_t': '2',
        'prm_podkozh_t': '10',
        'prm_retromam_t': '8',
        'prm_parenhima_t': '15',
        'levm_topografia': 'в норме',
        'levm_kozha_kontur': 'в норме',
        'levm_kozha_exogen': 'в норме',
        'levm_kozha_exostr': 'в норме',
        'levm_sosok': 'без особенностей',
        'levm_intramam': 'без особенностей',
        'levm_parenhima_kontur': 'в норме',
        'levm_parenhima_exogen': 'в норме',
        'levm_parenhima_exostr': 'без особенностей',
        'levm_galaktofori': 'не визуализируются',
        'levm_limfuzli': 'без особенностей',
        'levm_kozha_t': '2',
        'levm_podkozh_t': '10',
        'levm_retromam_t': '8',
        'levm_parenhima_t': '15',
        'zaklyuchenie': 'Молочные железы в норме',
        'rekomendacii': 'Ежегодное УЗИ'
    },
    8: {
        'mp_obem_fiz': '180',
        'mp_forma': 'в норме',
        'mp_kontur_naruzh': 'в норме',
        'mp_kontur_vnutr': 'в норме',
        'mp_stenka_do': '4',
        'mp_stenka_posle': '6',
        'mp_soderzhimoe': 'анэхогенное',
        'mp_sheika': 'в норме',
        'mp_ustya': 'в норме',
        'mp_ostatoch': '15',
        'zaklyuchenie': 'Мочевой пузырь в норме',
        'rekomendacii': 'Стандартное наблюдение'
    }
}

# ────────────────────────────────────────────────────────────────────────────
# SCENARIO 2: All pathological/abnormal options
# ────────────────────────────────────────────────────────────────────────────

SCENARIO_2_DATA = {
    1: {
        'pech_topografia': 'изменена',
        'pech_konturi': 'неровные',
        'pech_exostruktura': 'неоднородная',
        'pech_exogennost': 'повышенная',
        'pech_zvukoprov': 'снижена',
        'puz_topografia': 'смещён',
        'puz_forma': 'деформирована',
        'puz_kontur': 'нечёткий',
        'zhp_vnutrip': 'визуализируются',
        'puz_protok': 'расширен',
        'puz_soderzhimoe': 'с конкрементами',
        'pech_pzr_lev': '40',
        'pech_kkr_lev': '60',
        'pech_pzr_hvost': '15',
        'pech_pzr_prav': '80',
        'pech_kkr_prav': '100',
        'pech_kvr': '30',
        'pech_lat_razmer': '10',
        'pech_obem': '500',
        'pech_massa': '600',
        'soud_pech_veny': '15',
        'soud_sum_diam': '50',
        'soud_port_vena': '20',
        'soud_nizhnyaya_pv': '35',
        'zhp_obshchiy': '10',
        'puz_dlina': '30',
        'puz_tolshina': '5',
        'puz_shirina': '15',
        'puz_obem': '5',
        'puz_ploshad': '5',
        'puz_stenka': '5',
        'pech_ugol_prav': 'тупой (более 75°)',
        'pech_ugol_lev': 'тупой (более 45°)',
        'zaklyuchenie': 'Выявлены патологические изменения',
        'rekomendacii': 'Требуется дополнительное обследование'
    },
    2: {
        'prav_topografia': 'изменена',
        'prav_podvizhnost': 'ограничена',
        'prav_zhir_kapsul': 'с особенностями',
        'prav_fib_kapsul': 'с особенностями',
        'prav_konturi': 'нечёткие',
        'prav_exogen_par': 'повышена',
        'prav_exostr_par': 'с особенностями',
        'prav_chls_exostr': 'расширена',
        'prav_konkrement': 'визуализируются',
        'prav_mochetchnik': 'расширен',
        'prav_dlina': '150',
        'prav_tolshina': '30',
        'prav_shirina': '75',
        'prav_obem': '200',
        'prav_parenhima_t': '8',
        'prav_chls_tolshina': '10',
        'prav_kri': '0.5',
        'lev_topografia': 'смещена',
        'lev_podvizhnost': 'отсутствует',
        'lev_zhir_kapsul': 'с особенностями',
        'lev_fib_kapsul': 'с особенностями',
        'lev_konturi': 'неровные',
        'lev_exogen_par': 'снижена',
        'lev_exostr_par': 'с особенностями',
        'lev_chls_exostr': 'с конкрементами',
        'lev_konkrement': 'визуализируются',
        'lev_mochetchnik': 'не визуализируется',
        'lev_dlina': '70',
        'lev_tolshina': '25',
        'lev_shirina': '40',
        'lev_obem': '60',
        'lev_parenhima_t': '5',
        'lev_chls_tolshina': '15',
        'lev_kri': '0.2',
        'mp_obem': '300',
        'mp_forma': 'изменена',
        'mp_kontur_naruzh': 'неровный',
        'mp_kontur_vnutr': 'нечёткий',
        'mp_stenka_do': '10',
        'mp_stenka_posle': '12',
        'mp_soderzhimoe': 'с конкрементами',
        'mp_sheika': 'изменено',
        'mp_ustya': 'изменена',
        'mp_ostatoch': '100',
        'zaklyuchenie': 'Патология почек и мочевого пузыря'
    },
    3: {
        'p_topografiya': 'смещена',
        'p_kontur': 'нечёткий',
        'p_elastichnost': 'снижена',
        'p_podvijnost': 'ограничена',
        'p_razmer': '15',
        'p_ehostruktura': 'неоднородная',
        'p_ehogennost': 'повышенная',
        'dp_topografiya': 'изменена',
        'dp_kontur': 'неровный',
        'dp_elastichnost': 'снижена',
        'dp_podvijnost': 'ограничена',
        'dp_dlina': '80',
        'dp_tolshina': '45',
        'dp_shirina': '60',
        'dp_obem': '120',
        'dp_forma': 'деформирована',
        'dp_ehostruktura': 'диффузно неоднородная',
        'dp_ehogennost': 'гиперэхогенная',
        'dl_topografiya': 'смещена',
        'dl_kontur': 'нечёткий',
        'dl_elastichnost': 'снижена',
        'dl_podvijnost': 'ограничена',
        'dl_dlina': '75',
        'dl_tolshina': '40',
        'dl_shirina': '55',
        'dl_obem': '100',
        'dl_forma': 'изменена',
        'dl_ehostruktura': 'неоднородная',
        'dl_ehogennost': 'гипоэхогенная',
        'summarniy': '220',
        'kg': '100',
        'sm': '300',
        'zaklyuchenie_sh': 'Значительное увеличение щитовидной железы',
        'rekomendatsi_sh': 'Консультация эндокринолога, гормональное исследование',
        'vrach': 'А.Латипов',
        'telefon': '+992931234567'
    },
    4: {
        'pdj_topografia': 'смещена',
        'pdj_forma': 'деформирована',
        'pdj_kontur': 'неровный',
        'pdj_exostruktura': 'неоднородная',
        'pdj_exogennost': 'повышенная',
        'pdj_zvukoprov': 'снижена',
        'pdj_golovka': '50',
        'pdj_telo': '5',
        'pdj_hvost': '40',
        'pdj_dlina': '150',
        'pdj_virsungov_g': '8',
        'pdj_virsungov_t': '5',
        'zaklyuchenie': 'Признаки хронического панкреатита',
        'rekomendacii': 'Консультация гастроэнтеролога'
    },
    5: {
        'sel_topografia': 'смещена',
        'sel_forma': 'деформирована',
        'sel_konturi': 'нечёткие',
        'sel_exostruktura': 'неоднородная',
        'sel_exogennost': 'повышенная',
        'sel_zvukoprov': 'снижена',
        'sel_dlina': '200',
        'sel_tolshina': '80',
        'sel_shirina': '120',
        'sel_obem': '800',
        'sel_indeks': '50',
        'sel_lienalis': '12',
        'zaklyuchenie': 'Спленомегалия',
        'rekomendacii': 'Консультация гематолога'
    },
    6: {
        'pr_topografia': 'смещена',
        'pr_forma': 'деформирована',
        'pr_kontur': 'неровный',
        'pr_exostruktura': 'с кальцинатами',
        'pr_exogennost': 'повышена',
        'pr_urethra': 'изменена',
        'pr_kapsul_t': '3',
        'pr_dlina': '50',
        'pr_tolshina': '35',
        'pr_shirina': '50',
        'pr_obem': '70',
        'pr_indeks': '1.5',
        'pr_nach_obem': '250',
        'pr_ostatoch': '80',
        'svp_topografia': 'смещена',
        'svp_kontur': 'неровный',
        'svp_exostr_do': 'изменена',
        'svp_exostr_posle': 'изменена',
        'svp_exogennost': 'гипоэхогенная',
        'svp_dlina': '80',
        'svp_tolshina': '30',
        'svp_shirina': '40',
        'svl_topografia': 'смещена',
        'svl_kontur': 'неровный',
        'svl_exostr_do': 'изменена',
        'svl_exostr_posle': 'изменена',
        'svl_exogennost': 'гипоэхогенная',
        'svl_dlina': '75',
        'svl_tolshina': '28',
        'svl_shirina': '35',
        'zaklyuchenie': 'Аденома простаты с остаточной мочой',
        'rekomendacii': 'Консультация уролога'
    },
    7: {
        'prm_topografia': 'изменена',
        'prm_kozha_kontur': 'неровный',
        'prm_kozha_exogen': 'гиперэхогенная',
        'prm_kozha_exostr': 'неоднородная',
        'prm_sosok': 'с особенностями',
        'prm_intramam': 'с особенностями',
        'prm_parenhima_kontur': 'неровный',
        'prm_parenhima_exogen': 'гипоэхогенная',
        'prm_parenhima_exostr': 'с особенностями',
        'prm_galaktofori': 'визуализируются',
        'prm_limfuzli': 'с особенностями',
        'prm_kozha_t': '8',
        'prm_podkozh_t': '30',
        'prm_retromam_t': '25',
        'prm_parenhima_t': '45',
        'levm_topografia': 'смещена',
        'levm_kozha_kontur': 'нечёткий',
        'levm_kozha_exogen': 'гиперэхогенная',
        'levm_kozha_exostr': 'неоднородная',
        'levm_sosok': 'с особенностями',
        'levm_intramam': 'с особенностями',
        'levm_parenhima_kontur': 'нечёткий',
        'levm_parenhima_exogen': 'повышена',
        'levm_parenhima_exostr': 'с особенностями',
        'levm_galaktofori': 'визуализируются',
        'levm_limfuzli': 'с особенностями',
        'levm_kozha_t': '7',
        'levm_podkozh_t': '28',
        'levm_retromam_t': '20',
        'levm_parenhima_t': '40',
        'zaklyuchenie': 'Фиброкистозная мастопатия обеих молочных желез',
        'rekomendacii': 'Консультация маммолога'
    },
    8: {
        'mp_obem_fiz': '400',
        'mp_forma': 'деформирована',
        'mp_kontur_naruzh': 'неровный',
        'mp_kontur_vnutr': 'нечёткий',
        'mp_stenka_do': '15',
        'mp_stenka_posle': '20',
        'mp_soderzhimoe': 'с конкрементами',
        'mp_sheika': 'изменено',
        'mp_ustya': 'изменена',
        'mp_ostatoch': '200',
        'zaklyuchenie': 'Камни мочевого пузыря, постоянная остаточная моча',
        'rekomendacii': 'Консультация уролога для оперативного лечения'
    }
}

# ────────────────────────────────────────────────────────────────────────────
# SCENARIO 3: All fields empty
# ────────────────────────────────────────────────────────────────────────────

def create_empty_scenario(protocol_id):
    """Create empty field dict for given protocol"""
    fields = PROTOCOL_FIELDS_MAP[protocol_id]
    return {field: '' for field in fields}

# ────────────────────────────────────────────────────────────────────────────
# SCENARIO 4: Very long text in textarea fields
# ────────────────────────────────────────────────────────────────────────────

LONG_TEXT = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur."

SCENARIO_4_DATA = {
    1: {**create_empty_scenario(1), 'zaklyuchenie': LONG_TEXT, 'rekomendacii': LONG_TEXT},
    2: {**create_empty_scenario(2), 'zaklyuchenie': LONG_TEXT},
    3: {**create_empty_scenario(3), 'zaklyuchenie_sh': LONG_TEXT, 'rekomendatsi_sh': LONG_TEXT},
    4: {**create_empty_scenario(4), 'zaklyuchenie': LONG_TEXT, 'rekomendacii': LONG_TEXT},
    5: {**create_empty_scenario(5), 'zaklyuchenie': LONG_TEXT, 'rekomendacii': LONG_TEXT},
    6: {**create_empty_scenario(6), 'zaklyuchenie': LONG_TEXT, 'rekomendacii': LONG_TEXT},
    7: {**create_empty_scenario(7), 'zaklyuchenie': LONG_TEXT, 'rekomendacii': LONG_TEXT},
    8: {**create_empty_scenario(8), 'zaklyuchenie': LONG_TEXT, 'rekomendacii': LONG_TEXT},
}

# ────────────────────────────────────────────────────────────────────────────
# SCENARIO 5: Mixed scenario (some fields filled, some empty)
# ────────────────────────────────────────────────────────────────────────────

SCENARIO_5_DATA = {
    1: {
        'pech_topografia': 'в норме',
        'pech_konturi': '',
        'pech_exostruktura': 'однородная',
        'pech_exogennost': '',
        'puz_topografia': 'смещён',
        'puz_forma': '',
        'pech_pzr_lev': '50',
        'pech_kkr_lev': '',
        'puz_dlina': '70',
        'puz_tolshina': '',
        'zaklyuchenie': 'Частичная патология печени',
        'rekomendacii': ''
    },
    2: {
        'prav_topografia': 'в норме',
        'prav_podvizhnost': '',
        'prav_konturi': 'в норме',
        'prav_dlina': '100',
        'prav_tolshina': '',
        'lev_topografia': 'изменена',
        'lev_konturi': '',
        'lev_dlina': '90',
        'mp_obem': '200',
        'mp_forma': '',
        'zaklyuchenie': 'Ассиметрия почек'
    },
    3: {
        'p_topografiya': 'в норме',
        'p_kontur': '',
        'dp_topografiya': 'в норме',
        'dp_dlina': '50',
        'dp_tolshina': '',
        'dl_topografiya': '',
        'dl_dlina': '48',
        'summarniy': '',
        'kg': '70',
        'zaklyuchenie_sh': 'Щитовидная железа с элементами асимметрии'
    },
    4: {
        'pdj_topografia': 'в норме',
        'pdj_forma': '',
        'pdj_kontur': 'в норме',
        'pdj_golovka': '20',
        'pdj_telo': '',
        'pdj_hvost': '15',
        'zaklyuchenie': 'Поджелудочная железа'
    },
    5: {
        'sel_topografia': 'в норме',
        'sel_forma': '',
        'sel_konturi': 'в норме',
        'sel_dlina': '110',
        'sel_tolshina': '',
        'sel_obem': '150',
        'zaklyuchenie': 'Селезёнка в пределах нормы'
    },
    6: {
        'pr_topografia': 'в норме',
        'pr_forma': '',
        'pr_kontur': 'ровный, чёткий',
        'pr_dlina': '32',
        'pr_tolshina': '',
        'pr_obem': '22',
        'svp_topografia': 'в норме',
        'svp_dlina': '',
        'svl_topografia': '',
        'svl_dlina': '40',
        'zaklyuchenie': 'Простата'
    },
    7: {
        'prm_topografia': 'в норме',
        'prm_kozha_kontur': '',
        'prm_kozha_t': '2',
        'prm_parenhima_kontur': 'в норме',
        'prm_parenhima_t': '',
        'levm_topografia': 'в норме',
        'levm_kozha_kontur': '',
        'levm_parenhima_t': '15',
        'zaklyuchenie': 'Молочные железы'
    },
    8: {
        'mp_obem_fiz': '180',
        'mp_forma': '',
        'mp_kontur_naruzh': 'в норме',
        'mp_stenka_do': '',
        'mp_stenka_posle': '6',
        'mp_soderzhimoe': 'анэхогенное',
        'mp_ostatoch': '15',
        'zaklyuchenie': 'Мочевой пузырь'
    }
}

# ────────────────────────────────────────────────────────────────────────────
# TEST EXECUTION FUNCTIONS
# ────────────────────────────────────────────────────────────────────────────

def verify_docx_validity(docx_bytes):
    """Verify that the generated DOCX is valid and can be opened."""
    if not docx_bytes or len(docx_bytes) == 0:
        return False, "Generated DOCX is empty (0 bytes)"

    try:
        doc = Document(io.BytesIO(docx_bytes))
        text_content = '\n'.join([p.text for p in doc.paragraphs])
        return True, "Valid DOCX"
    except Exception as e:
        return False, f"Error opening DOCX: {str(e)}"

def check_fields_in_docx(docx_bytes, form_data):
    """Check if key fields appear in the generated DOCX."""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        all_text = '\n'.join([p.text for p in doc.paragraphs])

        # Check for key non-empty fields
        missing_fields = []
        found_fields = []

        for key, value in form_data.items():
            if value and len(str(value).strip()) > 0:
                # Search for value in document
                if str(value) in all_text:
                    found_fields.append(key)
                else:
                    missing_fields.append(f"{key}='{value}'")

        return len(missing_fields), found_fields, missing_fields
    except Exception as e:
        return -1, [], [f"Error checking fields: {str(e)}"]

def run_test_scenario(protocol_id, scenario_name, form_data):
    """Run a single test scenario for a protocol."""
    try:
        docx_bytes = generate_protocol_docx(
            protocol_id=protocol_id,
            protocol_title=f"Protocol {protocol_id}",
            form_data=form_data,
            client=CLIENT,
            doctor=DOCTOR,
            created_at=CREATED_AT
        )

        # Verify DOCX validity
        is_valid, validity_msg = verify_docx_validity(docx_bytes)

        if not is_valid:
            return {
                'protocol': protocol_id,
                'scenario': scenario_name,
                'success': False,
                'error': validity_msg,
                'bytes': 0,
                'missing_count': -1,
                'found_fields': [],
                'missing_fields': []
            }

        # Check fields
        missing_count, found_fields, missing_fields = check_fields_in_docx(docx_bytes, form_data)

        return {
            'protocol': protocol_id,
            'scenario': scenario_name,
            'success': True,
            'bytes': len(docx_bytes),
            'missing_count': missing_count,
            'found_fields': found_fields,
            'missing_fields': missing_fields[:5]  # Show first 5 missing
        }

    except Exception as e:
        return {
            'protocol': protocol_id,
            'scenario': scenario_name,
            'success': False,
            'error': f"Exception: {str(e)}",
            'bytes': 0,
            'missing_count': -1,
            'found_fields': [],
            'missing_fields': []
        }

def main():
    """Run comprehensive tests for all protocols."""
    results = {}
    total_tests = 0
    passed_tests = 0

    for protocol_id in range(1, 9):
        protocol_results = {}
        protocol_passed = 0
        protocol_total = 0

        print(f"\n{'='*80}")
        print(f"Protocol {protocol_id}")
        print(f"{'='*80}")

        # Scenario 1: Normal/healthy values
        scenario_1 = SCENARIO_1_DATA.get(protocol_id, {})
        result = run_test_scenario(protocol_id, "Scenario 1: Normal/Healthy", scenario_1)
        protocol_results[1] = result
        protocol_total += 1
        if result['success']:
            protocol_passed += 1
        print(f"  Scenario 1 (Normal): {'PASS' if result['success'] else 'FAIL'} - {result['bytes']} bytes")
        if not result['success']:
            print(f"    Error: {result.get('error', 'Unknown')}")

        # Scenario 2: All pathological/abnormal
        scenario_2 = SCENARIO_2_DATA.get(protocol_id, {})
        result = run_test_scenario(protocol_id, "Scenario 2: Pathological", scenario_2)
        protocol_results[2] = result
        protocol_total += 1
        if result['success']:
            protocol_passed += 1
        print(f"  Scenario 2 (Pathological): {'PASS' if result['success'] else 'FAIL'} - {result['bytes']} bytes")
        if not result['success']:
            print(f"    Error: {result.get('error', 'Unknown')}")

        # Scenario 3: All fields empty
        scenario_3 = create_empty_scenario(protocol_id)
        result = run_test_scenario(protocol_id, "Scenario 3: Empty Fields", scenario_3)
        protocol_results[3] = result
        protocol_total += 1
        if result['success']:
            protocol_passed += 1
        print(f"  Scenario 3 (Empty): {'PASS' if result['success'] else 'FAIL'} - {result['bytes']} bytes")
        if not result['success']:
            print(f"    Error: {result.get('error', 'Unknown')}")

        # Scenario 4: Very long text
        scenario_4 = SCENARIO_4_DATA.get(protocol_id, {})
        result = run_test_scenario(protocol_id, "Scenario 4: Long Text", scenario_4)
        protocol_results[4] = result
        protocol_total += 1
        if result['success']:
            protocol_passed += 1
        print(f"  Scenario 4 (Long Text): {'PASS' if result['success'] else 'FAIL'} - {result['bytes']} bytes")
        if not result['success']:
            print(f"    Error: {result.get('error', 'Unknown')}")

        # Scenario 5: Mixed
        scenario_5 = SCENARIO_5_DATA.get(protocol_id, {})
        result = run_test_scenario(protocol_id, "Scenario 5: Mixed", scenario_5)
        protocol_results[5] = result
        protocol_total += 1
        if result['success']:
            protocol_passed += 1
        print(f"  Scenario 5 (Mixed): {'PASS' if result['success'] else 'FAIL'} - {result['bytes']} bytes")
        if not result['success']:
            print(f"    Error: {result.get('error', 'Unknown')}")

        results[protocol_id] = {
            'passed': protocol_passed,
            'total': protocol_total,
            'scenarios': protocol_results
        }

        total_tests += protocol_total
        passed_tests += protocol_passed

    # Print summary
    print(f"\n{'='*80}")
    print("TEST SUMMARY")
    print(f"{'='*80}")

    for protocol_id in range(1, 9):
        proto_results = results[protocol_id]
        print(f"Protocol {protocol_id}: {proto_results['passed']}/{proto_results['total']} tests passed")

        # Show failures
        for scenario_num, result in proto_results['scenarios'].items():
            if not result['success']:
                print(f"  - Scenario {scenario_num} FAILED: {result.get('error', 'Unknown error')}")
            elif result.get('missing_count', 0) > 0:
                print(f"  - Scenario {scenario_num}: {result['missing_count']} fields missing")

    print(f"\n{'='*80}")
    print(f"TOTAL: {passed_tests}/{total_tests} tests passed ({100*passed_tests//total_tests}%)")
    print(f"{'='*80}\n")

    return passed_tests, total_tests

if __name__ == '__main__':
    passed, total = main()
    sys.exit(0 if passed == total else 1)
